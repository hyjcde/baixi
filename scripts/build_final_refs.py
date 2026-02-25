"""
最终版：直接用手工精修过的参考文献列表生成规范docx。
同时处理原docx的脚注格式规范化。
顺序：经史子集 → 今人专著 → 期刊论文 → 学位论文
各类内按年代排序。
脚注格式：去掉[M][J][D]等标记。
"""
import sys
import re
import copy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

# ═══════════════════════════════════════════
# 手工精修的参考文献列表
# ═══════════════════════════════════════════

REFS = {
    'jing': {
        'label': '（一）经部',
        'items': [
            '[汉]郑玄注：《周礼》，上海：商务印书馆，1936年。',
            '[汉]许慎撰，[清]段玉裁注，许惟贤整理：《说文解字注》，江苏：凤凰出版社，2007年。',
            '[汉]许慎撰，[宋]徐铉校定：《说文解字》，北京：中华书局，2015年。',
            '[汉]董仲舒撰，苏舆义证：《春秋繁露义证》，北京：中华书局，1992年。',
            '[汉]郑玄注，[唐]孔颖达疏：《礼记正义》，北京：北京大学出版社，1999年。',
            '[晋]杜预注：《春秋左传集解》，上海：上海人民出版社，1977年。',
            '[西晋]竺法护译：《佛说太子瑞应本起经》，载《大正藏》第3册，台北：新文丰出版公司，1983年。',
            '[北凉]昙无谶译：《大方等大集经》，载《大正藏》第13册，台北：新文丰出版公司，1983年。',
            '[清]孙诒让撰：《周礼正义》，北京：中华书局，1987年。',
            '[清]孙希旦撰，沈啸寰、王星贤点校：《礼记集解》，北京：中华书局，1989年。',
        ],
    },
    'shi': {
        'label': '（二）史部',
        'items': [
            '[汉]司马迁：《史记》，北京：中华书局，1959年。',
            '[汉]班固撰，[唐]颜师古注：《汉书》，北京：中华书局，1962年。',
            '[汉]刘向集录：《战国策》，上海：上海古籍出版社，1985年。',
            '[汉]刘向撰，张涛译注：《列女传译注》，济南：山东大学出版社，1990年。',
            '[汉]蔡质：《汉官典职仪式选用》，载[清]孙星衍辑《汉官六种》，北京：中华书局，1990年。',
            '[汉]刘歆等撰，王根林校点：《西京杂记》，上海：上海古籍出版社，2012年。',
            '[晋]陈寿撰，[南朝宋]裴松之注：《三国志》，北京：中华书局，1959年。',
            '[晋]葛洪撰，周天游校注：《西京杂记》，北京：中华书局，1985年。',
            '[晋]干宝撰，汪绍楹校注：《搜神记》，北京：中华书局，1979年。',
            '[晋]王嘉撰，[梁]萧绮录：《拾遗记》，北京：中华书局，1981年。',
            '[南朝宋]范晔撰，[唐]李贤等注：《后汉书》，北京：中华书局，1965年。',
            '[南朝梁]萧子显撰：《南齐书》，北京：中华书局，1972年。',
            '[南朝梁]沈约撰：《宋书》，北京：中华书局，1974年。',
            '[北魏]杨衒之撰，周祖谟校释：《洛阳伽蓝记校释》，北京：中华书局，2010年。',
            '[北齐]魏收撰：《魏书》，北京：中华书局，1974年。',
            '[唐]令狐德棻等撰：《周书》，北京：中华书局，1971年。',
            '[唐]李百药撰：《北齐书》，北京：中华书局，1972年。',
            '[唐]姚思廉等撰：《梁书》，北京：中华书局，1973年。',
            '[唐]魏徵等撰：《隋书》，北京：中华书局，1973年。',
            '[唐]房玄龄等撰：《晋书》，北京：中华书局，1974年。',
            '[唐]李延寿撰：《南史》，北京：中华书局，1975年。',
            '[唐]杜佑撰，王文锦等点校：《通典》，北京：中华书局，1988年。',
            '[唐]崔令钦撰，任半塘笺订：《教坊记笺订》，北京：中华书局，1962年。',
            '[唐]苏鹗：《杜阳杂编》，北京：中华书局，1958年。',
            '[唐]李绰：《尚书故实》，北京：中华书局，1958年。',
            '[唐]张鷟撰：《朝野佥载》，北京：中华书局，1979年。',
            '[唐]赵璘：《因话录》，北京：中华书局，1979年。',
            '[唐]段成式：《酉阳杂俎》，北京：中华书局，1981年。',
            '[唐]牛僧孺：《玄怪录》，北京：中华书局，1982年。',
            '[唐]李冗：《独异志》，北京：中华书局，1983年。',
            '[唐]孙頠：《幻异志·神女传》，北京：中华书局，1991年。',
            '[唐]郑处诲撰，田廷柱点校：《明皇杂录》，北京：中华书局，1994年。',
            '[唐]封演撰，赵贞信校注：《封氏闻见记校注》，北京：中华书局，2005年。',
            '[唐]姚汝能：《安禄山事迹》，北京：中华书局，2006年。',
            '[后晋]刘昫等撰：《旧唐书》，北京：中华书局，1975年。',
            '[五代]钱易：《南部新书》，北京：中华书局，2002年。',
            '[宋]司马光编著，[元]胡三省音注：《资治通鉴》，北京：中华书局，1956年。',
            '[宋]孟元老撰，邓之诚注：《东京梦华录注》，北京：中华书局，1982年。',
            '[宋]李昉等撰：《太平御览》，北京：中华书局，1960年。',
            '[宋]李昉等编：《太平广记》，北京：中华书局，1961年。',
            '[宋]王溥撰：《唐会要》，上海：上海古籍出版社，2006年。',
            '[北宋]欧阳修、宋祁撰：《新唐书》，北京：中华书局，1975年。',
            '[元]脱脱等撰：《宋史》，北京：中华书局，1977年。',
            '[元]马端临：《文献通考》，北京：中华书局，2011年。',
            '[明]胡震亨撰：《唐音癸签》，上海：上海古籍出版社，1981年。',
            '[清]高宗敕撰：《续通典》，北京：商务印书馆，1935年。',
            '[清]王应麟：《玉海》，上海：上海古籍出版社，1992年。',
        ],
    },
    'zi': {
        'label': '（三）子部',
        'items': [
            '[战国]列御寇撰，杨伯峻集释：《列子集释》，北京：中华书局，1979年。',
            '[汉]高诱注：《吕氏春秋》，上海：上海古籍出版社，2014年。',
            '[梁]释慧皎撰，汤用彤校注：《高僧传》，北京：中华书局，1992年。',
            '[唐]释道世撰：《法苑珠林》，北京：中华书局，2003年。',
            '[唐]郭象注，成玄英疏：《庄子注疏》，北京：中华书局，2011年。',
            '[宋]晁迥：《法藏碎金录》，明万历刻本。',
        ],
    },
    'ji': {
        'label': '（四）集部',
        'items': [
            '[南朝梁]萧统编，[唐]李善等注：《文选》，上海：上海古籍出版社，1986年。',
            '[南朝宋]鲍照：《舞鹤赋》，载萧统编《六臣注文选》卷十四，北京：中华书局，1987年。',
            '[三国魏]曹植著，赵幼文校注：《曹植集校注》，北京：人民文学出版社，1998年。',
            '[新罗]崔致远：《桂苑笔耕集》，首尔：韩国学中央研究院藏本。',
            '[唐]欧阳询撰，汪绍楹校：《艺文类聚》，上海：上海古籍出版社，1982年。',
            '[清]严可均辑：《全上古三代秦汉三国六朝文》，北京：中华书局，1958年。',
            '[清]彭定求等编：《全唐诗》，北京：中华书局，1960年。',
            '[清]彭定求等编：《全唐诗》（增订本），北京：中华书局，1999年。',
            '[清]董诰等编：《全唐文》，北京：中华书局，1983年。',
            '[清]逯钦立辑：《先秦汉魏晋南北朝诗》，北京：中华书局，1983年。',
        ],
    },
}

MODERN_BOOKS = [
    '王国维：《宋元戏曲史》，上海：商务印书馆，1915年。',
    '郭庆藩：《新编诸子集成·庄子集释》，北京：中华书局，1961年。',
    '周贻白：《中国戏曲发展史纲要》，上海：上海古籍出版社，1979年。',
    '袁珂校注：《山海经校注》，上海：上海古籍出版社，1980年。',
    '傅起凤、傅腾龙：《中国杂技》，天津：天津科学技术出版社，1983年。',
    '叶大兵：《中国百戏史话》，杭州：浙江人民出版社，1985年。',
    '陈永正主编：《中国方术大辞典》，广州：中山大学出版社，1991年。',
    '周楞伽：《幻术奇谈》，上海：上海古籍出版社，1993年。',
    '萧放：《〈荆楚岁时记〉研究：兼论传统中国民众生活的岁时观念》，北京：北京师范大学出版社，2000年。',
    '吉成名：《中国崇龙习俗研究》，天津：天津古籍出版社，2002年。',
    '胡大雷：《中古文学集团》，桂林：广西师范大学出版社，1996年。',
    '陈梦家：《商代的神话与巫术》，北京：中华书局，2006年。',
    '刘兴珍、李永林：《中华艺术通史·秦汉卷》，北京：北京师范大学出版社，2006年。',
    '郑传寅：《古代戏曲与东方文化》，武汉：武汉大学出版社，2007年。',
    '冀昀：《尚书》，北京：线装书局，2007年。',
    '蔡梦麒：《广韵校释》，长沙：岳麓书院，2007年。',
    '方勇译注：《庄子》，北京：中华书局，2010年。',
]

JOURNALS = [
    '刘永连：《舞马和马舞》，《中国文化研究》2005年第3期。',
    '黄水云：《汉代游艺赋初探》，《中国楚辞学》2009年第2期。',
    '钱志熙：《南北朝隋代散乐与戏剧关系札论》，《文学与文化》2010年。',
]

THESES = [
    # 目前从脚注中未发现学位论文
]


def build_docx(output_path):
    """生成参考文献docx"""
    doc = Document()

    # 默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    rpr = style.element.get_or_add_rPr()
    rpr.set(qn('w:rFonts'), '宋体')

    def add_text(text, bold=False, size=Pt(12)):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = '宋体'
        run.font.size = size
        run.bold = bold
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        return p

    # 大标题
    h = doc.add_heading('参考文献', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.name = '黑体'
        run.font.size = Pt(22)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 一、古籍文献
    add_text('一、古籍文献', bold=True, size=Pt(14))

    for cat in ['jing', 'shi', 'zi', 'ji']:
        section = REFS[cat]
        add_text(f'    {section["label"]}', bold=True, size=Pt(13))
        for i, item in enumerate(section['items'], 1):
            add_text(f'{i}. {item}')

    # 二、今人专著
    add_text('二、今人专著', bold=True, size=Pt(14))
    for i, item in enumerate(MODERN_BOOKS, 1):
        add_text(f'{i}. {item}')

    # 三、期刊论文
    add_text('三、期刊论文', bold=True, size=Pt(14))
    for i, item in enumerate(JOURNALS, 1):
        add_text(f'{i}. {item}')

    # 四、学位论文
    if THESES:
        add_text('四、学位论文', bold=True, size=Pt(14))
        for i, item in enumerate(THESES, 1):
            add_text(f'{i}. {item}')

    doc.save(output_path)
    print(f"参考文献已保存到: {output_path}")


def normalize_footnotes(docx_path, output_path):
    """规范化脚注格式：去掉[M][J][D]等标记"""
    doc = Document(docx_path)
    
    footnotes_part = None
    for rel in doc.part.rels.values():
        if "footnotes" in rel.reltype:
            footnotes_part = rel.target_part
            break
    
    if not footnotes_part:
        print("未找到脚注")
        return
    
    root = etree.fromstring(footnotes_part.blob)
    modified = False
    
    for fn in root.findall('.//w:footnote', NSMAP):
        fn_id = fn.get(f'{{{NSMAP["w"]}}}id')
        if fn_id in ('0', '-1'):
            continue
        
        for t in fn.iter(f'{{{NSMAP["w"]}}}t'):
            if t.text:
                original = t.text
                # 去掉文献类型标记
                new_text = re.sub(r'\[M\d*\]', '', original)
                new_text = re.sub(r'［Ｍ］', '', new_text)
                new_text = re.sub(r'\[([JDCAGZN])\]', '', new_text)
                # 全角句点改逗号
                new_text = new_text.replace('．', '，')
                
                if new_text != original:
                    t.text = new_text
                    modified = True
    
    if modified:
        footnotes_part._blob = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
        doc.save(output_path)
        print(f"已规范化脚注并保存到: {output_path}")
    else:
        print("脚注中未发现需要修改的格式标记")


def print_refs():
    """打印参考文献预览"""
    print("=" * 70)
    print("参  考  文  献")
    print("=" * 70)

    print("\n一、古籍文献\n")
    for cat in ['jing', 'shi', 'zi', 'ji']:
        section = REFS[cat]
        print(f"\n    {section['label']}\n")
        for i, item in enumerate(section['items'], 1):
            print(f"  {i}. {item}")

    print("\n\n二、今人专著\n")
    for i, item in enumerate(MODERN_BOOKS, 1):
        print(f"  {i}. {item}")

    print("\n\n三、期刊论文\n")
    for i, item in enumerate(JOURNALS, 1):
        print(f"  {i}. {item}")

    if THESES:
        print("\n\n四、学位论文\n")
        for i, item in enumerate(THESES, 1):
            print(f"  {i}. {item}")

    total = sum(len(REFS[c]['items']) for c in REFS) + len(MODERN_BOOKS) + len(JOURNALS) + len(THESES)
    print(f"\n\n总计 {total} 条参考文献")


if __name__ == '__main__':
    docx_path = sys.argv[1] if len(sys.argv) > 1 else '郑璎_《汉唐百戏书写研究》二稿参考文献1_20260225.docx'
    
    # 1. 打印预览
    print_refs()
    
    # 2. 生成参考文献docx
    ref_output = docx_path.replace('.docx', '_参考文献_final.docx')
    build_docx(ref_output)
    
    # 3. 规范化脚注
    fn_output = docx_path.replace('.docx', '_脚注规范化.docx')
    normalize_footnotes(docx_path, fn_output)
