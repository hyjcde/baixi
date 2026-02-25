"""
v2: 精细处理参考文献。
  - 载于总集（《全唐诗》《全唐文》《文选》《艺文类聚》等）的单篇不单列，只列总集
  - 同一书不同版本/注本合并（保留最完整的）
  - 缺失作者的手工补全
  - 赋/诗正确归入集部
  - 输出到docx
"""
import sys, re, copy
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

DYNASTY_ORDER = {
    '原题': 0, '春秋': 1, '战国': 2, '秦': 3,
    '汉': 4, '西汉': 4, '东汉': 5,
    '三国': 6, '三国魏': 6, '蜀汉': 6, '魏': 6, '吴': 6,
    '西晋': 7, '东晋': 7, '晋': 7,
    '十六国': 8, '北凉': 8, '前秦': 8,
    '南朝宋': 9, '南朝齐': 10, '南朝梁': 11, '南朝陈': 12,
    '南朝': 10, '梁': 11, '陈': 12,
    '北魏': 13, '东魏': 14, '西魏': 14, '北齐': 15, '北周': 15,
    '隋': 16, '唐': 17, '五代': 18, '后晋': 18, '后唐': 18,
    '新罗': 17,
    '宋': 19, '北宋': 19, '南宋': 20,
    '辽': 19, '金': 20, '元': 21, '明': 22, '清': 23,
}

# ────────── 手工条目：补全脚本无法自动识别/解析的文献 ──────────
# 格式: (dynasty, author, book, city, publisher, year, category, type)
MANUAL_ENTRIES = [
    # 经部
    ('晋', '杜预注', '春秋左传集解', '上海', '上海人民出版社', '1977', 'jing', 'ancient'),
    ('清', '孙诒让撰', '周礼正义', '北京', '中华书局', '1987', 'jing', 'ancient'),
    ('清', '孙希旦撰，沈啸寰、王星贤点校', '礼记集解', '北京', '中华书局', '1989', 'jing', 'ancient'),
    # 史部
    ('宋', '孟元老撰', '东京梦华录', '北京', '古典文学出版社', '1957', 'shi', 'ancient'),
    # 子部
    ('汉', '高诱注', '吕氏春秋', '上海', '上海古籍出版社', '2014', 'zi', 'ancient'),
    ('唐', '郭象注，成玄英疏', '庄子注疏', '北京', '中华书局', '2011', 'zi', 'ancient'),
    # 今人专著
    ('', '王国维', '宋元戏曲史', '上海', '商务印书馆', '1915', '', 'modern_book'),
    ('', '陈梦家', '商代的神话与巫术', '北京', '中华书局', '2006', '', 'modern_book'),
    ('', '叶大兵', '中国百戏史话', '杭州', '浙江人民出版社', '1985', '', 'modern_book'),
    ('', '周贻白', '中国戏曲发展史纲要', '上海', '上海古籍出版社', '1979', '', 'modern_book'),
    ('', '郑传寅', '古代戏曲与东方文化', '武汉', '武汉大学出版社', '2007', '', 'modern_book'),
    ('', '刘兴珍、李永林', '中华艺术通史·秦汉卷', '北京', '北京师范大学出版社', '2006', '', 'modern_book'),
    ('', '吉成名', '中国崇龙习俗研究', '天津', '天津古籍出版社', '2002', '', 'modern_book'),
    # 期刊论文
    ('', '钱志熙', '南北朝隋代散乐与戏剧关系札论', '', '《文学与文化》', '2010', '', 'journal'),
    ('', '黄水云', '汉代游艺赋初探', '', '《中国楚辞学》', '2009', '', 'journal'),
    ('', '刘永连', '舞马和马舞', '', '《中国文化研究》', '2005', '', 'journal'),
]

# 载于总集的单篇 → 参考文献只列总集，不列单篇
ANTHOLOGY_BOOKS = {
    '全唐诗', '全唐文', '文选', '六臣注文选',
    '全上古三代秦汉三国六朝文', '先秦汉魏晋南北朝诗',
    '艺文类聚', '太平御览', '太平广记',
    '全后汉文', '全晋文', '全三国文', '全宋文',
}


def extract_footnotes(docx_path):
    doc = Document(docx_path)
    footnotes_part = None
    for rel in doc.part.rels.values():
        if "footnotes" in rel.reltype:
            footnotes_part = rel.target_part
            break
    if not footnotes_part:
        return []
    root = etree.fromstring(footnotes_part.blob)
    results = []
    for fn in root.findall('.//w:footnote', NSMAP):
        fn_id = fn.get(f'{{{NSMAP["w"]}}}id')
        if fn_id in ('0', '-1'):
            continue
        texts = [t.text for t in fn.iter(f'{{{NSMAP["w"]}}}t') if t.text]
        full_text = ''.join(texts).strip()
        if full_text:
            results.append((fn_id, full_text))
    return results


def normalize_text(text):
    text = re.sub(r'\[M\d*\]', '', text)
    text = re.sub(r'［Ｍ］', '', text)
    text = re.sub(r'\[([JDCAGZN])\]', '', text)
    text = text.replace('．', '，')
    return text.strip()


def parse_ref(text):
    info = {'original': text, 'dynasty': '', 'author': '', 'book': '',
            'annotator': '', 'publisher': '', 'year': '', 'city': '',
            'volume': '', 'page': '', 'series': '', 'type': 'unknown',
            'hosted_in': ''}

    # [朝代]作者
    m = re.match(r'[\[［【]([^\]］】]+)[\]］】](.+?)(?=(?:：|:|《|，))', text)
    if m:
        info['dynasty'] = m.group(1).strip()
        info['author'] = m.group(2).strip()
    else:
        m2 = re.match(r'^([^：:《\[，]{1,20})(?:：|:)', text)
        if m2:
            info['author'] = m2.group(1).strip()

    # 书名
    books = re.findall(r'《([^》]+)》', text)
    if books:
        info['book'] = books[0]
        # 检查是否 "载于" 某总集
        if len(books) >= 2:
            for anthology in ANTHOLOGY_BOOKS:
                for b in books[1:]:
                    if anthology in b:
                        info['hosted_in'] = anthology
                        break

    # 城市
    for c in ['北京', '上海', '天津', '重庆', '南京', '杭州', '广州', '武汉',
              '成都', '西安', '长沙', '沈阳', '济南', '台北', '首尔', '桂林',
              '江苏', '中国台北']:
        if c in text:
            info['city'] = c
            break

    # 出版社
    pm = re.search(r'([\u4e00-\u9fa5]+(?:出版[\u4e00-\u9fa5]*|书局|书院|印书馆|出版公司))', text)
    if pm:
        info['publisher'] = pm.group(1)

    # 年份
    ym = re.search(r'(\d{4})\s*年', text)
    if ym:
        info['year'] = ym.group(1)

    # 页码
    pgm = re.search(r'第\s*([^页]+?)\s*页', text)
    if pgm:
        info['page'] = pgm.group(1).strip()

    # 类型
    if info['dynasty'] and info['dynasty'] in DYNASTY_ORDER:
        info['type'] = 'ancient'
    elif re.search(r'第\s*\d+\s*期', text):
        info['type'] = 'journal'
    elif '学位' in text or '硕士' in text or '博士' in text:
        info['type'] = 'thesis'
    elif info['author'] and not info['dynasty']:
        info['type'] = 'modern_book'

    return info


def classify_ancient(book, text=''):
    """经史子集分类"""
    jing = ['周礼', '仪礼', '礼记', '春秋左传', '春秋繁露', '左传',
            '周易', '尚书', '毛诗', '论语', '孟子', '尔雅', '孝经',
            '十三经', '礼记正义', '周礼正义',
            '说文解字', '风俗通义',
            '大正藏', '大方等大集经', '佛说太子瑞应本起经']
    shi = ['史记', '汉书', '后汉书', '三国志', '晋书', '宋书', '南齐书',
           '梁书', '陈书', '魏书', '北齐书', '周书', '隋书', '南史', '北史',
           '旧唐书', '新唐书', '宋史',
           '资治通鉴', '通典', '文献通考', '续通典',
           '唐会要', '玉海',
           '西京杂记', '洛阳伽蓝记', '邺中记',
           '安禄山事迹', '明皇杂录', '因话录', '封氏闻见记',
           '朝野佥载', '独异志', '南部新书', '教坊记', '尚书故实', '杜阳杂编',
           '东京梦华录', '梦粱录', '中朝故事',
           '唐音癸签', '战国策', '列女传',
           '荆楚岁时记', '岁时广记', '古今岁时杂咏', '玉烛宝典', '四民月令',
           '汉官典职', '汉官六种', '睡虎地秦墓竹简',
           '太平广记', '太平御览',
           '酉阳杂俎', '幻异志', '幻戏志', '玄怪录',
           '搜神记', '拾遗记', '册府元龟']
    zi = ['庄子', '老子', '韩非子', '荀子', '墨子',
          '管子', '淮南子', '吕氏春秋', '列子',
          '山海经', '广韵',
          '高僧传', '法苑珠林', '法藏碎金录']
    ji = ['文选', '六臣注文选',
          '全唐诗', '全唐文', '全上古三代秦汉三国六朝文',
          '先秦汉魏晋南北朝诗',
          '艺文类聚', '初学记',
          '诗品', '曹植集', '桂苑笔耕集']

    for kw in jing:
        if kw in book: return 'jing'
    for kw in shi:
        if kw in book: return 'shi'
    for kw in zi:
        if kw in book: return 'zi'
    for kw in ji:
        if kw in book: return 'ji'
    return 'shi'


def make_entry(dynasty, author, book, city, publisher, year):
    """构建一条标准条目字典"""
    return {
        'dynasty': dynasty, 'author': author, 'book': book,
        'city': city, 'publisher': publisher, 'year': year,
        'original': '',
    }


def format_ref(info):
    """格式化为参考文献条目"""
    parts = []
    if info.get('dynasty'):
        parts.append(f"[{info['dynasty']}]")
    if info.get('author'):
        parts.append(info['author'])
    if info.get('book'):
        book_clean = info['book']
        book_clean = re.sub(r'卷[^》，]*', '', book_clean).strip()
        if '·' in book_clean and not any(k in book_clean for k in ['通史', '集成']):
            book_clean = re.sub(r'·[^》]*', '', book_clean).strip()
        parts.append(f"：《{book_clean}》")
    if info.get('city') and info.get('publisher'):
        parts.append(f"，{info['city']}：{info['publisher']}")
    elif info.get('publisher'):
        parts.append(f"，{info['publisher']}")
    if info.get('year'):
        parts.append(f"，{info['year']}年。")
    else:
        parts.append("。")
    result = ''.join(parts).replace('，，', '，').replace('。。', '。')
    return result


def format_journal(info):
    parts = []
    if info.get('author'):
        parts.append(info['author'])
    if info.get('book'):
        parts.append(f"：《{info['book']}》")
    if info.get('publisher'):
        pub = info['publisher']
        if not pub.startswith('《'):
            pub = f"《{pub}》"
        parts.append(f"，{pub}")
    if info.get('year'):
        parts.append(f"{info['year']}年")
    pm = re.search(r'第\s*(\d+)\s*期', info.get('original', ''))
    if pm:
        parts.append(f"第{pm.group(1)}期。")
    else:
        parts.append("。")
    return ''.join(parts)


def dynasty_key(info):
    d = info.get('dynasty', '')
    order = DYNASTY_ORDER.get(d, 99)
    y = int(info['year']) if info.get('year', '').isdigit() else 9999
    return (order, y, info.get('author', ''))


def book_key(info):
    """用于去重的键"""
    bk = info.get('book', '')
    bk = re.sub(r'卷.*', '', bk).strip()
    bk = re.sub(r'·.*', '', bk).strip()
    bk = re.sub(r'[（(].*[）)]', '', bk).strip()
    return bk


def dedup(items):
    seen = OrderedDict()
    for it in items:
        key = book_key(it)
        if not key:
            continue
        if key in seen:
            old = seen[key]
            if len(it.get('original', '')) > len(old.get('original', '')):
                seen[key] = it
            elif not old.get('author') and it.get('author'):
                seen[key] = it
        else:
            seen[key] = it
    return list(seen.values())


# ──────────────────────────────────
# 主流程
# ──────────────────────────────────
def main(docx_path):
    footnotes = extract_footnotes(docx_path)
    print(f"共 {len(footnotes)} 条脚注\n")

    # 收集所有引用到的总集
    anthologies_needed = set()

    ancient = {'jing': [], 'shi': [], 'zi': [], 'ji': []}
    modern_books = []
    journals = []
    theses = []
    skipped = []

    skip_ids = set()  # 手工条目覆盖的脚注ID

    # 先加入手工条目
    for dynasty, author, book, city, publisher, year, category, typ in MANUAL_ENTRIES:
        entry = make_entry(dynasty, author, book, city, publisher, year)
        entry['type'] = typ
        if typ == 'ancient' and category:
            ancient[category].append(entry)
        elif typ == 'modern_book':
            modern_books.append(entry)
        elif typ == 'journal':
            entry['publisher'] = publisher  # 期刊名
            journals.append(entry)

    for fn_id, raw in footnotes:
        text = normalize_text(raw)

        # 跳过
        if text.startswith('同上') or text.strip() == '？' or text.startswith('？'):
            skipped.append((fn_id, text[:60]))
            continue
        if '《' not in text and '出版' not in text and '书局' not in text:
            skipped.append((fn_id, text[:60]))
            continue

        info = parse_ref(text)
        if not info['book']:
            skipped.append((fn_id, text[:60]))
            continue

        # 如果是载于总集的单篇 → 记录总集，跳过单篇
        if info['hosted_in']:
            anthologies_needed.add(info['hosted_in'])
            # 但仍需记录总集本身（后面去重处理）

        # 对于出处为总集的诗文赋，只列总集
        # 检查是否为单篇（书名是篇名而非总集名）
        bk = info['book']
        is_single_piece = False
        if info['hosted_in']:
            is_single_piece = True
        elif any(bk.endswith(suf) for suf in ['赋', '歌', '行', '诗', '伎']):
            # 可能是单篇
            for anth in ANTHOLOGY_BOOKS:
                if anth in text:
                    is_single_piece = True
                    anthologies_needed.add(anth)
                    break

        if is_single_piece:
            # 解析总集信息
            for anth_name in ANTHOLOGY_BOOKS:
                if anth_name in text:
                    # 提取总集的编者信息
                    anth_match = re.search(rf'([\[［][^\]］]+[\]］][^：:]+?)(?:：|:).*?《{re.escape(anth_name)}》', text)
                    if anth_match:
                        anth_info = parse_ref(text)
                        anth_info['book'] = anth_name
                        # 保留总集的编者
                        cat = classify_ancient(anth_name, text)
                        ancient[cat].append(anth_info)
            continue

        info['fn_id'] = fn_id

        if info['type'] == 'ancient':
            cat = classify_ancient(info['book'], text)
            ancient[cat].append(info)
        elif info['type'] == 'journal':
            journals.append(info)
        elif info['type'] == 'thesis':
            theses.append(info)
        elif info['type'] == 'modern_book':
            modern_books.append(info)
        else:
            if '出版' in text or '书局' in text or '年' in text:
                cat = classify_ancient(info['book'], text)
                ancient[cat].append(info)
            else:
                skipped.append((fn_id, text[:60]))

    # 去重+排序
    for cat in ancient:
        ancient[cat] = dedup(ancient[cat])
        ancient[cat].sort(key=dynasty_key)
    modern_books = dedup(modern_books)
    modern_books.sort(key=lambda x: int(x.get('year', '9999')) if x.get('year', '').isdigit() else 9999)
    journals = dedup(journals)
    journals.sort(key=lambda x: int(x.get('year', '9999')) if x.get('year', '').isdigit() else 9999)

    # ── 打印 ──
    print("=" * 70)
    print("参  考  文  献")
    print("=" * 70)

    print("\n一、古籍文献")
    cat_labels = [('jing','（一）经部'), ('shi','（二）史部'), ('zi','（三）子部'), ('ji','（四）集部')]
    for cat, label in cat_labels:
        items = ancient[cat]
        if not items:
            continue
        print(f"\n    {label}\n")
        for i, it in enumerate(items, 1):
            print(f"  {i}. {format_ref(it)}")

    print("\n\n二、今人专著\n")
    for i, it in enumerate(modern_books, 1):
        print(f"  {i}. {format_ref(it)}")

    print("\n\n三、期刊论文\n")
    for i, it in enumerate(journals, 1):
        print(f"  {i}. {format_journal(it)}")

    if theses:
        print("\n\n四、学位论文\n")
        for i, it in enumerate(theses, 1):
            print(f"  {i}. {format_ref(it)}")

    print(f"\n\n已跳过 {len(skipped)} 条:")
    for fid, t in skipped:
        print(f"  [{fid}] {t}")

    # ── 输出 docx ──
    out_path = docx_path.replace('.docx', '_参考文献_v2.docx')
    build_docx(ancient, modern_books, journals, theses, out_path)
    return ancient, modern_books, journals, theses


def build_docx(ancient, modern_books, journals, theses, out_path):
    doc = Document()

    # 设置正文默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    h = doc.add_heading('参考文献', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('一、古籍文献', level=2)
    for cat, label in [('jing','（一）经部'), ('shi','（二）史部'), ('zi','（三）子部'), ('ji','（四）集部')]:
        items = ancient.get(cat, [])
        if not items:
            continue
        doc.add_heading(f'    {label}', level=3)
        for i, it in enumerate(items, 1):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(f"{i}. {format_ref(it)}")
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    if modern_books:
        doc.add_heading('二、今人专著', level=2)
        for i, it in enumerate(modern_books, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {format_ref(it)}")
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    if journals:
        doc.add_heading('三、期刊论文', level=2)
        for i, it in enumerate(journals, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {format_journal(it)}")
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    if theses:
        doc.add_heading('四、学位论文', level=2)
        for i, it in enumerate(theses, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{i}. {format_ref(it)}")
            run.font.name = '宋体'
            run.font.size = Pt(12)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.save(out_path)
    print(f"\n参考文献已保存到: {out_path}")


if __name__ == '__main__':
    path = sys.argv[1] if len(sys.argv) > 1 else '郑璎_《汉唐百戏书写研究》二稿参考文献1_20260225.docx'
    main(path)
