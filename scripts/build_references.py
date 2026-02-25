"""
从论文docx提取脚注，生成规范化参考文献列表+修正脚注格式，输出为新docx。

参考文献格式（师姐格式，无[M][J][D]）：
  古籍：[朝代]作者撰：《书名》，城市：出版社，年份年。
  今人专著：作者：《书名》，城市：出版社，年份年。
  期刊论文：作者：《篇名》，《刊名》年份年第X期。

参考文献顺序：经史子集 → 今人专著 → 期刊论文 → 学位论文
各类内按年代排序。
"""
import sys
import re
import copy
from collections import OrderedDict
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

DYNASTY_ORDER = {
    '原题': 0,
    '春秋': 1, '战国': 2, '秦': 3,
    '汉': 4, '西汉': 4, '东汉': 5,
    '三国': 6, '三国魏': 6, '蜀汉': 6, '魏': 6, '吴': 6,
    '西晋': 7, '东晋': 7, '晋': 7,
    '十六国': 8, '北凉': 8, '前秦': 8,
    '南朝宋': 9, '南朝齐': 10, '南朝梁': 11, '南朝陈': 12,
    '南朝': 10, '梁': 11, '陈': 12,
    '北魏': 13, '东魏': 14, '西魏': 14, '北齐': 15, '北周': 15,
    '隋': 16, '唐': 17, '五代': 18, '后晋': 18, '后唐': 18,
    '新罗': 17,
    '宋': 19, '北宋': 19, '南宋': 20,
    '辽': 19, '金': 20, '元': 21, '明': 22, '清': 23,
}

# ──────────── 手工修正表 ────────────
# key = 脚注编号(str), value = 规范后的参考文献条目（用于参考文献列表，不含页码）
MANUAL_REFS = {
    # 被跳过的但包含引用的脚注
    '110': {'dynasty': '晋', 'author': '杜预注', 'book': '春秋左传集解', 'city': '上海', 'publisher': '上海人民出版社', 'year': '1977', 'category': 'jing', 'type': 'ancient'},
    '123': {'dynasty': '', 'author': '冀昀', 'book': '尚书', 'city': '北京', 'publisher': '线装书局', 'year': '2007', 'category': 'jing', 'type': 'modern_book'},
    '125': {'dynasty': '', 'author': '王国维', 'book': '宋元戏曲史', 'city': '', 'publisher': '', 'year': '', 'category': 'modern_book', 'type': 'modern_book'},
    '126': {'dynasty': '清', 'author': '孙诒让撰', 'book': '周礼正义', 'city': '北京', 'publisher': '中华书局', 'year': '1987', 'category': 'jing', 'type': 'ancient'},
    '127': {'dynasty': '清', 'author': '孙希旦撰，沈啸寰、王星贤点校', 'book': '礼记集解', 'city': '北京', 'publisher': '中华书局', 'year': '1989', 'category': 'jing', 'type': 'ancient'},
    '136': {'dynasty': '', 'author': '刘永连', 'book': '舞马和马舞', 'city': '', 'publisher': '《中国文化研究》', 'year': '2005', 'category': 'journal', 'type': 'journal'},
    '54': {'dynasty': '', 'author': '黄水云', 'book': '汉代游艺赋初探', 'city': '', 'publisher': '《中国楚辞学》', 'year': '2009', 'category': 'journal', 'type': 'journal'},
    '128': {'dynasty': '', 'author': '陈梦家', 'book': '商代的神话与巫术', 'city': '', 'publisher': '', 'year': '', 'category': 'modern_book', 'type': 'modern_book'},
}


def extract_footnotes(docx_path):
    doc = Document(docx_path)
    footnotes_part = None
    for rel in doc.part.rels.values():
        if "footnotes" in rel.reltype:
            footnotes_part = rel.target_part
            break
    if not footnotes_part:
        return []
    root = etree.fromstring(footnotes_part.blob)
    results = []
    for fn in root.findall('.//w:footnote', NSMAP):
        fn_id = fn.get(f'{{{NSMAP["w"]}}}id')
        if fn_id in ('0', '-1'):
            continue
        texts = [t.text for t in fn.iter(f'{{{NSMAP["w"]}}}t') if t.text]
        full_text = ''.join(texts).strip()
        if full_text:
            results.append((fn_id, full_text))
    return results


def normalize_text(text):
    """去掉[M][J]等标记，统一标点"""
    text = re.sub(r'\[M\d*\]', '', text)
    text = re.sub(r'\[([JDCAGZN])\]', '', text)
    text = re.sub(r'\[Ｍ\]', '', text)
    text = re.sub(r'［Ｍ］', '', text)
    text = text.replace('．', '，')
    # 不要全局去空格，保留中文间有意义的空格
    return text.strip()


def parse_single_ref(text):
    """解析一条引用为结构化信息"""
    info = {'original': text, 'dynasty': '', 'author': '', 'book': '',
            'annotator': '', 'publisher': '', 'year': '', 'city': '',
            'volume': '', 'page': '', 'series': '', 'type': 'unknown'}

    # [朝代]作者
    m = re.match(r'[\[［【]([^\]］】]+)[\]］】](.+?)(?:：|:)', text)
    if m:
        info['dynasty'] = m.group(1).strip()
        info['author'] = m.group(2).strip()
    else:
        m2 = re.match(r'^([^：:《\[，]{1,15})(?:：|:)', text)
        if m2:
            info['author'] = m2.group(1).strip()

    # 书名《...》
    books = re.findall(r'《([^》]+)》', text)
    if books:
        info['book'] = books[0]
        if len(books) > 1:
            info['series'] = books[1] if len(books) > 1 else ''

    # 卷
    vm = re.search(r'卷([^，,。：:]+)', text)
    if vm:
        info['volume'] = vm.group(1).strip()

    # 城市
    cities = ['北京', '上海', '天津', '重庆', '南京', '杭州', '广州', '武汉',
              '成都', '西安', '长沙', '沈阳', '济南', '台北', '首尔', '桂林',
              '长春', '哈尔滨', '石家庄', '郑州', '合肥', '福州', '南昌',
              '太原', '兰州', '昆明', '贵阳', '海口', '银川', '西宁', '江苏',
              '中国台北']
    for c in cities:
        if c in text:
            info['city'] = c
            break

    # 出版社
    pm = re.search(r'([\u4e00-\u9fa5]+(?:出版[\u4e00-\u9fa5]*|书局|书院|印书馆))', text)
    if pm:
        info['publisher'] = pm.group(1)

    # 年份
    ym = re.search(r'(\d{4})\s*年', text)
    if ym:
        info['year'] = ym.group(1)

    # 页码
    pgm = re.search(r'第\s*([^页]+?)\s*页', text)
    if pgm:
        info['page'] = pgm.group(1).strip()

    # 类型判断
    if info['dynasty'] and info['dynasty'] in DYNASTY_ORDER:
        info['type'] = 'ancient'
    elif re.search(r'第\s*\d+\s*期', text) or re.search(r'\d{4}\s*年\s*第\s*\d+\s*期', text) or '期刊' in text:
        info['type'] = 'journal'
    elif '学位' in text or '硕士' in text or '博士' in text:
        info['type'] = 'thesis'
    elif info['author'] and not info['dynasty']:
        info['type'] = 'modern_book'

    return info


def classify_ancient(book, text=''):
    """经史子集分类"""
    jing = ['周礼', '仪礼', '礼记', '春秋', '左传', '公羊', '穀梁',
            '周易', '易经', '尚书', '毛诗', '诗经', '论语', '孟子', '尔雅', '孝经',
            '十三经', '礼记正义', '春秋左传', '周礼正义',
            '春秋繁露', '说文解字',
            '大正藏', '大方等大集经', '佛说太子瑞应本起经',
            '风俗通义']
    shi = ['史记', '汉书', '后汉书', '三国志', '晋书', '宋书', '南齐书',
           '梁书', '陈书', '魏书', '北齐书', '周书', '隋书', '南史', '北史',
           '旧唐书', '新唐书', '宋史', '资治通鉴', '通典', '文献通考', '续通典',
           '唐会要', '玉海',
           '西京杂记', '洛阳伽蓝记', '邺中记',
           '安禄山事迹', '明皇杂录', '因话录', '封氏闻见记',
           '朝野佥载', '独异志', '南部新书', '教坊记', '尚书故实', '杜阳杂编',
           '东京梦华录', '梦粱录', '中朝故事',
           '唐音癸签', '战国策', '列女传',
           '荆楚岁时记', '岁时广记', '古今岁时杂咏', '玉烛宝典', '四民月令',
           '汉官典职', '汉官六种', '睡虎地秦墓竹简',
           '太平广记', '太平御览',
           '酉阳杂俎', '幻异志', '幻戏志', '玄怪录',
           '搜神记', '拾遗记',
           '册府元龟']
    zi = ['庄子', '老子', '韩非子', '荀子', '墨子',
          '管子', '淮南子', '吕氏春秋', '列子',
          '广韵', '论衡', '颜氏家训',
          '新编诸子集成', '诸子集成',
          '山海经',
          '高僧传', '法苑珠林', '法藏碎金录',
          '中国方术大辞典', '幻术奇谈']
    ji = ['文选', '六臣注文选', '文心雕龙',
          '全唐诗', '全唐文', '全上古三代秦汉三国六朝文',
          '先秦汉魏晋南北朝诗',
          '艺文类聚', '初学记', '古文苑',
          '诗品', '曹植集', '玉台新咏',
          '桂苑笔耕集']

    for kw in jing:
        if kw in book: return 'jing'
    for kw in shi:
        if kw in book: return 'shi'
    for kw in zi:
        if kw in book: return 'zi'
    for kw in ji:
        if kw in book: return 'ji'

    # 赋、诗、歌 → 集部（但要排除诗经等）
    if re.search(r'赋[》]|[》].*赋$', book) or '桂苑' in book:
        return 'ji'

    return 'shi'


def format_ref_entry(info):
    """格式化为参考文献条目（不含页码和卷数）"""
    parts = []
    if info.get('dynasty'):
        parts.append(f"[{info['dynasty']}]")
    if info.get('author'):
        parts.append(info['author'])

    book = info.get('book', '')
    # 参考文献列表中去掉卷数和篇名
    book_clean = re.sub(r'卷[^》，]*', '', book).strip()
    book_clean = re.sub(r'·[^》]*', '', book_clean).strip() if '·' in book_clean else book_clean
    if book_clean:
        parts.append(f"：《{book_clean}》")

    if info.get('city') and info.get('publisher'):
        parts.append(f"，{info['city']}：{info['publisher']}")
    elif info.get('publisher'):
        parts.append(f"，{info['publisher']}")

    if info.get('year'):
        parts.append(f"，{info['year']}年。")
    else:
        parts.append("。")

    result = ''.join(parts)
    result = result.replace('，，', '，').replace('。。', '。')
    return result


def format_journal_entry(info):
    """格式化期刊论文条目"""
    parts = []
    if info.get('author'):
        parts.append(info['author'])
    if info.get('book'):
        parts.append(f"：《{info['book']}》")
    if info.get('series'):
        parts.append(f"，《{info['series']}》")
    if info.get('year'):
        parts.append(f"{info['year']}年")
    # 从原文提取期数
    pm = re.search(r'第\s*(\d+)\s*期', info.get('original', ''))
    if pm:
        parts.append(f"第{pm.group(1)}期。")
    else:
        parts.append("。")
    return ''.join(parts)


def dynasty_key(info):
    d = info.get('dynasty', '')
    order = DYNASTY_ORDER.get(d, 99)
    y = int(info['year']) if info.get('year', '').isdigit() else 9999
    return (order, y, info.get('author', ''))


def dedup(items):
    """按 (book_key, author) 去重"""
    seen = OrderedDict()
    for it in items:
        bk = it.get('book', '')
        bk_key = re.sub(r'卷.*', '', bk).strip()
        bk_key = re.sub(r'·.*', '', bk_key).strip()
        bk_key = re.sub(r'[（(].*[）)]', '', bk_key).strip()
        author_key = it.get('author', '')[:4]
        key = (bk_key, author_key)
        if key not in seen or len(it.get('original', '')) > len(seen[key].get('original', '')):
            seen[key] = it
    return list(seen.values())


# ──────────────────────────────────────
# 写参考文献到 docx
# ──────────────────────────────────────
def build_refs_docx(ancient, modern_books, journals, theses, output_path):
    """生成仅包含参考文献的docx（之后可追加到原文后面）"""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 标题
    h = doc.add_heading('参考文献', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 一、古籍文献
    doc.add_heading('一、古籍文献', level=2)

    cat_labels = OrderedDict([
        ('jing', '（一）经部'),
        ('shi', '（二）史部'),
        ('zi', '（三）子部'),
        ('ji', '（四）集部'),
    ])

    for cat, label in cat_labels.items():
        items = ancient.get(cat, [])
        if not items:
            continue
        doc.add_heading(label, level=3)
        for i, it in enumerate(items, 1):
            entry = format_ref_entry(it)
            doc.add_paragraph(f"{i}. {entry}", style='List Number')

    # 二、今人专著
    if modern_books:
        doc.add_heading('二、今人专著', level=2)
        for i, it in enumerate(modern_books, 1):
            entry = format_ref_entry(it)
            doc.add_paragraph(f"{i}. {entry}", style='List Number')

    # 三、期刊论文
    if journals:
        doc.add_heading('三、期刊论文', level=2)
        for i, it in enumerate(journals, 1):
            entry = format_journal_entry(it)
            doc.add_paragraph(f"{i}. {entry}", style='List Number')

    # 四、学位论文
    if theses:
        doc.add_heading('四、学位论文', level=2)
        for i, it in enumerate(theses, 1):
            entry = format_ref_entry(it)
            doc.add_paragraph(f"{i}. {entry}", style='List Number')

    doc.save(output_path)
    print(f"\n参考文献已保存到: {output_path}")


# ──────────────────────────────────────
# 主流程
# ──────────────────────────────────────
def main(docx_path):
    footnotes = extract_footnotes(docx_path)
    print(f"共提取 {len(footnotes)} 条脚注\n")

    ancient = {'jing': [], 'shi': [], 'zi': [], 'ji': []}
    modern_books = []
    journals = []
    theses = []
    skipped = []

    for fn_id, raw_text in footnotes:
        text = normalize_text(raw_text)

        # 手工修正
        if fn_id in MANUAL_REFS:
            m = MANUAL_REFS[fn_id]
            cat = m.get('category', '')
            if m['type'] == 'ancient' and cat in ancient:
                ancient[cat].append(m)
            elif m['type'] == 'journal':
                journals.append(m)
            elif m['type'] == 'modern_book':
                modern_books.append(m)
            continue

        # 跳过"同上"、"？"、纯说明性注释
        if text.startswith('同上') or text == '？' or text.startswith('？'):
            skipped.append((fn_id, text[:60]))
            continue
        if '《' not in text:
            skipped.append((fn_id, text[:60]))
            continue

        # 解析
        info = parse_single_ref(text)
        if not info['book']:
            skipped.append((fn_id, text[:60]))
            continue

        info['fn_id'] = fn_id

        if info['type'] == 'ancient':
            cat = classify_ancient(info['book'], text)
            ancient[cat].append(info)
        elif info['type'] == 'journal':
            journals.append(info)
        elif info['type'] == 'thesis':
            theses.append(info)
        elif info['type'] == 'modern_book':
            modern_books.append(info)
        else:
            # 有《》但无法判断类型 → 尝试按古籍处理
            if '出版' in text or '书局' in text or '年' in text:
                cat = classify_ancient(info['book'], text)
                ancient[cat].append(info)
            else:
                skipped.append((fn_id, text[:60]))

    # 去重 + 排序
    for cat in ancient:
        ancient[cat] = dedup(ancient[cat])
        ancient[cat].sort(key=dynasty_key)

    modern_books = dedup(modern_books)
    modern_books.sort(key=lambda x: int(x.get('year', '9999')) if x.get('year', '').isdigit() else 9999)

    journals = dedup(journals)
    journals.sort(key=lambda x: int(x.get('year', '9999')) if x.get('year', '').isdigit() else 9999)

    # 打印预览
    print("=" * 70)
    print("参  考  文  献")
    print("=" * 70)

    print("\n一、古籍文献")
    for cat, label in [('jing','（一）经部'), ('shi','（二）史部'), ('zi','（三）子部'), ('ji','（四）集部')]:
        print(f"\n    {label}\n")
        for i, it in enumerate(ancient[cat], 1):
            print(f"  {i}. {format_ref_entry(it)}")

    print("\n\n二、今人专著\n")
    for i, it in enumerate(modern_books, 1):
        print(f"  {i}. {format_ref_entry(it)}")

    print("\n\n三、期刊论文\n")
    for i, it in enumerate(journals, 1):
        print(f"  {i}. {format_journal_entry(it)}")

    if theses:
        print("\n\n四、学位论文\n")
        for i, it in enumerate(theses, 1):
            print(f"  {i}. {format_ref_entry(it)}")

    print(f"\n\n已跳过 {len(skipped)} 条非引用/无法解析的脚注:")
    for fn_id, t in skipped:
        print(f"  [{fn_id}] {t}")

    # 生成参考文献 docx
    out_path = docx_path.replace('.docx', '_参考文献.docx')
    build_refs_docx(ancient, modern_books, journals, theses, out_path)

    return ancient, modern_books, journals, theses, skipped


if __name__ == '__main__':
    path = sys.argv[1] if len(sys.argv) > 1 else '郑璎_《汉唐百戏书写研究》二稿参考文献1_20260225.docx'
    main(path)
