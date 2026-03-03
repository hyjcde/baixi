"""
为《汉唐百戏书写研究》添加页眉页脚。

页眉规范（参考论文格式规范.pdf）：
- 宋体五号，居中
- 奇数页页眉为章节标题，偶数页页眉为论文题目（或统一居中章节标题）
- 页眉下方有横线

页脚规范：
- 页码居中，宋体小五号
- 摘要/目录部分使用罗马数字页码
- 正文部分使用阿拉伯数字页码（从绪论开始重新编号为1）

各章节页眉内容：
- 摘要  -> "摘要"
- 目录  -> "目录"
- 绪论  -> "绪论"
- 第一章 -> "第一章 汉唐百戏书写概述"
- 第二章 -> "第二章 杂技类百戏书写研究"
- 第三章 -> "第三章 驯兽与拟兽类百戏书写研究"
- 第四章 -> "第四章 幻术类百戏书写研究"
- 结语  -> "结语"
- 附录1 -> "附录"
- 附录2 -> "附录"
- 参考文献 -> "参考文献"
"""

import copy
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree


INPUT_FILE = "../郑璎_《汉唐百戏书写研究》_20260303.docx"
OUTPUT_FILE = "../郑璎_《汉唐百戏书写研究》_20260303.docx"

SECTION_DEFS = [
    {"para_idx": 0,   "header_text": "摘要",                     "page_num_fmt": "upperRoman", "restart_page": True,  "start_num": 1},
    {"para_idx": 11,  "header_text": "目录",                     "page_num_fmt": "upperRoman", "restart_page": False, "start_num": None},
    {"para_idx": 37,  "header_text": "绪论",                     "page_num_fmt": "decimal",    "restart_page": True,  "start_num": 1},
    {"para_idx": 105, "header_text": "第一章 汉唐百戏书写概述",    "page_num_fmt": "decimal",    "restart_page": False, "start_num": None},
    {"para_idx": 194, "header_text": "第二章 杂技类百戏书写研究",  "page_num_fmt": "decimal",    "restart_page": False, "start_num": None},
    {"para_idx": 371, "header_text": "第三章 驯兽与拟兽类百戏书写研究", "page_num_fmt": "decimal", "restart_page": False, "start_num": None},
    {"para_idx": 598, "header_text": "第四章 幻术类百戏书写研究",  "page_num_fmt": "decimal",    "restart_page": False, "start_num": None},
    {"para_idx": 748, "header_text": "结语",                     "page_num_fmt": "decimal",    "restart_page": False, "start_num": None},
    {"para_idx": 755, "header_text": "附录",                     "page_num_fmt": "decimal",    "restart_page": False, "start_num": None},
    {"para_idx": 758, "header_text": "附录",                     "page_num_fmt": "decimal",    "restart_page": False, "start_num": None},
    {"para_idx": 761, "header_text": "参考文献",                  "page_num_fmt": "decimal",    "restart_page": False, "start_num": None},
]


def get_section_properties_template(doc):
    """Extract the final sectPr from the document body as a template for page size/margins."""
    body = doc.element.body
    sect_pr = body.find(qn('w:sectPr'))
    if sect_pr is not None:
        return copy.deepcopy(sect_pr)
    return None


def create_sect_pr(template_sect_pr, section_type="nextPage"):
    """Create a new w:sectPr element copying page size and margins from template."""
    new_sect_pr = copy.deepcopy(template_sect_pr)

    for child in list(new_sect_pr):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag not in ('pgSz', 'pgMar', 'cols'):
            new_sect_pr.remove(child)

    type_elem = parse_xml(f'<w:type {nsdecls("w")} w:val="{section_type}"/>')
    new_sect_pr.insert(0, type_elem)

    return new_sect_pr


def set_header_content(header, text):
    """Set header paragraph with 宋体五号 centered text and bottom border."""
    header.is_linked_to_previous = False

    for p in header.paragraphs:
        p.clear()

    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(text)
    run.font.size = Pt(10.5)  # 五号 = 10.5pt
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/>'
        f'</w:pBdr>'
    )
    pBdr_existing = pPr.find(qn('w:pBdr'))
    if pBdr_existing is not None:
        pPr.remove(pBdr_existing)
    pPr.append(pBdr)


def set_footer_with_page_number(footer, fmt="decimal"):
    """Set footer with centered page number field."""
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p.clear()

    if footer.paragraphs:
        p = footer.paragraphs[0]
    else:
        p = footer.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_begin = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/><w:szCs w:val="18"/><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/></w:rPr></w:r>')
    run_begin.append(fld_char_begin)

    instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_instr = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/><w:szCs w:val="18"/><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/></w:rPr></w:r>')
    run_instr.append(instr_text)

    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_end = parse_xml(f'<w:r {nsdecls("w")}><w:rPr><w:sz w:val="18"/><w:szCs w:val="18"/><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/></w:rPr></w:r>')
    run_end.append(fld_char_end)

    p._element.append(run_begin)
    p._element.append(run_instr)
    p._element.append(run_end)


def set_page_number_format(sect_pr, fmt="decimal", start=None):
    """Set the page number format on the section properties.
    fmt: 'decimal', 'upperRoman', 'lowerRoman'
    start: if not None, restart page numbering at this value.
    """
    pgNumType = sect_pr.find(qn('w:pgNumType'))
    if pgNumType is not None:
        sect_pr.remove(pgNumType)

    attrs = f'w:fmt="{fmt}"'
    if start is not None:
        attrs += f' w:start="{start}"'
    pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")} {attrs}/>')
    sect_pr.append(pgNumType)


def add_header_reference(sect_pr, rel_id):
    """Add w:headerReference to sectPr."""
    ref = parse_xml(
        f'<w:headerReference {nsdecls("w", "r")} w:type="default" r:id="{rel_id}"/>'
    )
    sect_pr.append(ref)


def add_footer_reference(sect_pr, rel_id):
    """Add w:footerReference to sectPr."""
    ref = parse_xml(
        f'<w:footerReference {nsdecls("w", "r")} w:type="default" r:id="{rel_id}"/>'
    )
    sect_pr.append(ref)


def main():
    doc = Document(INPUT_FILE)
    template_sect_pr = get_section_properties_template(doc)

    if template_sect_pr is None:
        print("ERROR: Cannot find section properties template in document.")
        return

    sorted_defs = sorted(SECTION_DEFS, key=lambda x: x["para_idx"], reverse=True)

    for i, sec_def in enumerate(sorted_defs):
        para_idx = sec_def["para_idx"]
        para = doc.paragraphs[para_idx]
        para_elem = para._element

        if para_idx == 0:
            continue

        new_sect_pr = create_sect_pr(template_sect_pr, "nextPage")
        set_page_number_format(new_sect_pr, sec_def["page_num_fmt"], sec_def.get("start_num"))

        prev_para_idx = para_idx - 1
        prev_para = doc.paragraphs[prev_para_idx]
        prev_pPr = prev_para._element.get_or_add_pPr()

        existing_sect = prev_pPr.find(qn('w:sectPr'))
        if existing_sect is not None:
            prev_pPr.remove(existing_sect)

        prev_pPr.append(new_sect_pr)

    final_sect_pr = doc.element.body.find(qn('w:sectPr'))
    if final_sect_pr is not None:
        last_def = SECTION_DEFS[-1]
        set_page_number_format(final_sect_pr, last_def["page_num_fmt"], last_def.get("start_num"))

    doc.save(OUTPUT_FILE)
    print(f"Saved with section breaks to {OUTPUT_FILE}")

    doc = Document(OUTPUT_FILE)
    print(f"\nDocument now has {len(doc.sections)} sections")

    section_headers = []
    for sec_def in SECTION_DEFS:
        section_headers.append(sec_def["header_text"])

    if len(section_headers) < len(doc.sections):
        first_def = SECTION_DEFS[0]
        section_headers = [first_def["header_text"]] + section_headers

    for i, section in enumerate(doc.sections):
        if i < len(SECTION_DEFS):
            sec_def = SECTION_DEFS[i]
        else:
            sec_def = SECTION_DEFS[-1]

        header_text = sec_def["header_text"]
        page_fmt = sec_def["page_num_fmt"]

        header = section.header
        set_header_content(header, header_text)

        footer = section.footer
        set_footer_with_page_number(footer, page_fmt)

        sect_pr = section._sectPr
        set_page_number_format(sect_pr, sec_def["page_num_fmt"], sec_def.get("start_num"))

        print(f"Section {i}: header='{header_text}', page_fmt={page_fmt}, restart={sec_def.get('start_num')}")

    doc.save(OUTPUT_FILE)
    print(f"\nFinal document saved to {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
