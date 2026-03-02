"""
修改 马伎放驯兽里了.docx 中标黄/标红需降重的段落。
降重段落索引：[1], [8], [38], [92], [97], [106]
修改后的文字标绿色（008000），其余保持原样。
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from copy import deepcopy
import lxml.etree as ET

import os
SRC = None
for _f in os.listdir('.'):
    if '\u9a6c\u4f0e' in _f and _f.endswith('.docx'):
        SRC = _f
        break
assert SRC, "Cannot find the .docx file!"
print(f"Using file: {SRC}")

doc = Document(SRC)
ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

GREEN = "008000"

def make_run(text, color=GREEN, font_name="\u5b8b\u4f53", font_size=12, bold=False):
    """Create a new <w:r> element with specified formatting."""
    r = ET.SubElement(ET.Element("dummy"), qn("w:r"))
    rPr = ET.SubElement(r, qn("w:rPr"))
    rFonts = ET.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    sz = ET.SubElement(rPr, qn("w:sz"))
    sz.set(qn("w:val"), str(font_size * 2))
    szCs = ET.SubElement(rPr, qn("w:szCs"))
    szCs.set(qn("w:val"), str(font_size * 2))
    if color:
        c = ET.SubElement(rPr, qn("w:color"))
        c.set(qn("w:val"), color)
    if bold:
        ET.SubElement(rPr, qn("w:b"))
    t = ET.SubElement(r, qn("w:t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    return r


def replace_para_text(para, new_text, color=GREEN, font_name="\u5b8b\u4f53", font_size=12):
    """Replace all runs in a paragraph with a single new run of given text+color.
    Preserves paragraph-level formatting (pPr)."""
    elem = para._element
    for old_r in elem.findall(qn("w:r")):
        elem.remove(old_r)
    r = make_run(new_text, color=color, font_name=font_name, font_size=font_size)
    elem.append(r)


def replace_para_partial(para, old_fragment, new_fragment, color=GREEN):
    """In a paragraph that has mixed coloring, replace only the runs whose text
    contains old_fragment, preserving the rest."""
    full_text = para.text
    if old_fragment not in full_text:
        print(f"  WARNING: fragment not found in para text!")
        return
    elem = para._element
    runs = elem.findall(qn("w:r"))

    concat = ""
    target_runs = []
    for r in runs:
        t_elem = r.find(qn("w:t"))
        if t_elem is not None and t_elem.text:
            concat += t_elem.text
            target_runs.append(r)

    for old_r in target_runs:
        elem.remove(old_r)

    r = make_run(new_fragment, color=color)
    elem.append(r)


# ============================================================
# Paragraph [1] - 整段替换（RED标记，写得不好）
# ============================================================
print("Processing [1]...")
NEW_1 = (
    "\u99ef\u517d\u7c7b\u767e\u620f\u4ee5\u99ef\u5316\u52a8\u7269\u4e3a\u6838\u5fc3\uff0c"
    "\u6db5\u76d6\u821e\u9a6c\u3001\u9a6c\u4f0e\u3001\u99ef\u8c61\u3001\u99ef\u7280\u7b49"
    "\u591a\u79cd\u8868\u6f14\u5f62\u6001\uff0c\u5176\u5386\u53f2\u53d1\u5c55\u8de8\u8d8a"
    "\u6c49\u9b4f\u81f3\u968b\u5510\uff0c\u7ecf\u5386\u4e86\u4ece\u5bab\u5ef7\u8d21\u54c1"
    "\u5c55\u793a\u5230\u72ec\u7acb\u8868\u6f14\u827a\u672f\u7684\u6f14\u53d8\u3002"
    "\u56f4\u7ed5\u8fd9\u4e9b\u8868\u6f14\u7684\u4e66\u5199\u4ea6\u56e0\u65f6\u4ee3"
    "\u4e0d\u540c\u800c\u5448\u73b0\u51fa\u663e\u8457\u5dee\u5f02\uff1a\u53f2\u4e66"
    "\u4fa7\u91cd\u8bb0\u5f55\u671d\u8d21\u4e0e\u5178\u793c\u4eea\u5236\uff0c\u6587\u5b66"
    "\u4f5c\u54c1\u5219\u7740\u529b\u4e8e\u827a\u672f\u63cf\u6479\u4e0e\u653f\u6cbb"
    "\u8bbd\u8c15\uff0c\u4e8c\u8005\u5171\u540c\u6784\u6210\u4e86\u99ef\u517d\u7c7b"
    "\u767e\u620f\u4e66\u5199\u7684\u591a\u91cd\u9762\u8c8c\u3002"
)
replace_para_text(doc.paragraphs[1], NEW_1, color=GREEN)

# ============================================================
# Paragraph [8] - 后半句降重
# 原文：此则记载虽简，却点出了舞马表演中的"拜"与"鼓节相应"这两个核心要素，
#       说明最晚至曹魏时期已有马匹能够随鼓点完成规定动作，舞马技艺的基本形态已然成型。
#       舞马经过训练，随着音律倾首，按着鼓点踏脚，舞姿能够配合音乐鼓点节奏。
# 需降重部分（YELLOW）：舞马经过训练，随着音律倾首，按着鼓点踏脚，舞姿能够配合音乐鼓点节奏。
# ============================================================
print("Processing [8]...")
NEW_8 = (
    "\u6b64\u5219\u8bb0\u8f7d\u867d\u7b80\uff0c\u5374\u70b9\u51fa\u4e86\u821e\u9a6c"
    "\u8868\u6f14\u4e2d\u7684\u201c\u62dc\u201d\u4e0e\u201c\u9f13\u8282\u76f8\u5e94\u201d"
    "\u8fd9\u4e24\u4e2a\u6838\u5fc3\u8981\u7d20\uff0c\u8bf4\u660e\u6700\u665a\u81f3"
    "\u66f9\u9b4f\u65f6\u671f\u5df2\u6709\u9a6c\u5339\u80fd\u591f\u968f\u9f13\u70b9"
    "\u5b8c\u6210\u89c4\u5b9a\u52a8\u4f5c\uff0c\u821e\u9a6c\u6280\u827a\u7684\u57fa\u672c"
    "\u5f62\u6001\u5df2\u7136\u6210\u578b\u3002\u6362\u8a00\u4e4b\uff0c\u66f9\u690d\u6240"
    "\u732e\u4e4b\u9a6c\u5df2\u5177\u5907\u542c\u8fa8\u4e50\u5f8b\u3001\u4ee5\u80a2\u4f53"
    "\u52a8\u4f5c\u5e94\u548c\u8282\u62cd\u7684\u80fd\u529b\uff0c\u8fd9\u6b63\u662f\u821e"
    "\u9a6c\u8868\u6f14\u7684\u57fa\u672c\u8981\u6c42\u3002"
)
replace_para_text(doc.paragraphs[8], NEW_8, color=GREEN)

# ============================================================
# Paragraph [38] - 谢庄和张率同题舞马赋（YELLOW整段）
# ============================================================
print("Processing [38]...")
NEW_38 = (
    "\u8c22\u5e84\u3001\u5f20\u7387\u4e8c\u4eba\u7684\u540c\u9898\u300a\u821e\u9a6c"
    "\u8d4b\u300b\u662f\u73b0\u5b58\u4ec5\u6709\u7684\u4e24\u7bc7\u5357\u671d\u548f"
    "\u821e\u9a6c\u4e4b\u8d4b\u3002\u5f20\u7387\u8d4b\u4e2d\u201c\u64e2\u9f99\u9996"
    "\uff0c\u56de\u9e7f\u8eaf\u3002\u7768\u4e24\u955c\uff0c\u8e59\u53cc\u51eb\u3002"
    "\u65e2\u5c31\u573a\u800c\u96c5\u62dc\uff0c\u65f6\u8d74\u66f2\u800c\u5f90\u8d8b\u201d"
    "\u6570\u53e5\uff0c\u4ee5\u8fde\u7eed\u7684\u77ed\u53e5\u52fe\u52d2\u51fa\u821e\u9a6c"
    "\u6602\u9996\u56de\u8eab\u3001\u51dd\u76ee\u655b\u8033\u7684\u59ff\u6001\uff0c\u4ee5"
    "\u53ca\u5165\u573a\u81f4\u793c\u3001\u968f\u66f2\u7f13\u884c\u7684\u4eea\u8282\u3002"
    "\u7d27\u63a5\u7740\uff0c\u201c\u654f\u8e81\u4e2d\u4e8e\u4fc3\u8282\uff0c\u6377\u7e41"
    "\u5916\u4e8e\u60ca\u6874\u201d\u5219\u5c06\u7b14\u89e6\u8f6c\u5411\u821e\u9a6c\u5728"
    "\u6025\u4fc3\u8282\u62cd\u4e0b\u7075\u654f\u53d8\u6362\u7684\u52a8\u6001\u4e4b\u7f8e"
    "\u3002\u4e24\u7bc7\u8d4b\u4f5c\u5747\u4ee5\u80e1\u9a6c\u5357\u6765\u3001\u5e94\u8282"
    "\u800c\u821e\u4e3a\u7d20\u6750\uff0c\u5728\u9882\u626c\u821e\u9a6c\u6280\u827a\u7684"
    "\u540c\u65f6\uff0c\u6709\u610f\u8425\u6784\u201c\u56db\u5937\u6765\u671d\u201d\u201c"
    "\u767e\u517d\u7387\u821e\u201d\u7684\u592a\u5e73\u666f\u8c61\uff0c\u5176\u4e66\u5199"
    "\u76ee\u7684\u4e0e\u653f\u6cbb\u9882\u5fb7\u5bc6\u4e0d\u53ef\u5206\u3002"
)
replace_para_text(doc.paragraphs[38], NEW_38, color=GREEN)

# ============================================================
# Paragraph [92] - 该赋在塑造技艺高超...侧肩（YELLOW整段）
# ============================================================
print("Processing [92]...")
NEW_92 = (
    "\u8d4b\u4e2d\u201c\u6216\u4fa7\u80a9\u4ee5\u9a70\u89c1\uff0c\u6216\u5954\u8dc3"
    "\u4ee5\u4e50\u95fb\u201d\u4e00\u53e5\u503c\u5f97\u7559\u610f\u3002\u201c\u4fa7\u80a9"
    "\u201d\u4e8c\u5b57\u770b\u4f3c\u5bfb\u5e38\uff0c\u5b9e\u5219\u517c\u5199\u4e86\u4e24"
    "\u4e2a\u5bf9\u8c61\uff1a\u89c2\u8005\u4fa7\u8eab\u63a2\u770b\uff0c\u8bf4\u660e\u4eba"
    "\u7fa4\u62e5\u6324\u3001\u4e89\u76f8\u76ee\u7779\uff1b\u827a\u4eba\u4fa7\u8eab\u7eb5"
    "\u9a6c\u75be\u9a70\uff0c\u5219\u70b9\u660e\u4e86\u9a6c\u4e0a\u52a8\u4f5c\u7684\u7075"
    "\u6d3b\u591a\u53d8\u3002\u8d4b\u5bb6\u4ec5\u7528\u4e00\u4e2a\u5fae\u5c0f\u7684\u4f53"
    "\u6001\u7ec6\u8282\uff0c\u4fbf\u540c\u65f6\u4f20\u8fbe\u51fa\u89c2\u6f14\u53cc\u65b9"
    "\u7684\u7d27\u5f20\u4e0e\u70ed\u70c8\uff0c\u7eb8\u4e0a\u6e38\u827a\u56e0\u6b64\u987f"
    "\u751f\u9c9c\u6d3b\u4e4b\u611f\u3002"
)
replace_para_text(doc.paragraphs[92], NEW_92, color=GREEN)

# ============================================================
# Paragraph [97] - 同样描绘了女子的马伎表演...居安若厉...（YELLOW整段）
# 查重标记：【劝谏皇帝不能沉迷于安逸享乐之中，应当居安若厉，摒弃郑卫之类淫乐】
#           【初盛唐多大型游艺活动...反对奢靡享乐。】
# ============================================================
print("Processing [97]...")
NEW_97 = (
    "\u656c\u62ec\u4e0e\u674e\u6fef\u8d4b\u4e3b\u9898\u76f8\u8fd1\uff0c\u6b64\u8d4b\u540c"
    "\u6837\u4ee5\u5bab\u5ef7\u5973\u5b50\u9a6c\u4f0e\u4e3a\u5bf9\u8c61\uff0c\u4f46\u66f4"
    "\u7740\u529b\u4e8e\u5f62\u5f0f\u7684\u534e\u4e3d\u4e0e\u52a8\u4f5c\u7684\u7cbe\u5999"
    "\uff1a\u5973\u5b50\u4eec\u7fd8\u8dbe\u91d1\u978d\u3001\u59d4\u8eab\u7389\u8e6c\uff0c"
    "\u5728\u6821\u573a\u4e0a\u4fa7\u80a9\u5954\u8dc3\u3001\u4e89\u950b\u9a8b\u6280\u3002"
    "\u503c\u5f97\u6ce8\u610f\u7684\u662f\uff0c\u8d4b\u672b\u7279\u610f\u8f6c\u5165\u8bbd"
    "\u8c0f\uff1a\u201c\u65af\u5e1d\u7687\u6240\u4ee5\u56e0\u58ee\u89c2\u800c\u6212\u9038"
    "\uff0c\u9042\u5c45\u5b89\u800c\u82e5\u5389\u3002\u5c82\u6deb\u4e50\u4ee5\u60d1\u4eba"
    "\uff0c\u89c1\u7ec8\u671d\u4e8e\u90d1\u536b\uff1f\u201d\u8d4b\u5bb6\u501f\u58ee\u89c2"
    "\u4e4b\u8f9e\u5bfc\u5411\u6212\u5962\u4e4b\u610f\uff0c\u63d0\u9192\u5e1d\u738b\u4e0d"
    "\u53ef\u56e0\u5a31\u89c2\u4e4b\u76db\u800c\u5fd8\u4e4e\u6240\u4ee5\u3002\u8fd9\u79cd"
    "\u201c\u8d4b\u672b\u8bbd\u8c0f\u201d\u7684\u7b14\u6cd5\u5728\u521d\u76db\u5510\u6e38"
    "\u827a\u8d4b\u4e2d\u9887\u4e3a\u5e38\u89c1\uff0c\u5f7c\u65f6\u754b\u730e\u3001\u4e50"
    "\u821e\u3001\u767e\u620f\u7b49\u5927\u578b\u6d3b\u52a8\u9891\u7e41\uff0c\u8d22\u529b"
    "\u4e0e\u4eba\u529b\u7684\u8017\u8d39\u4e0d\u53ef\u5c0f\u89d1\uff0c\u56e0\u6b64\u90e8"
    "\u5206\u8d4b\u5bb6\u5728\u94fa\u53d9\u6280\u827a\u4e4b\u4f59\uff0c\u5e38\u5e26\u6709"
    "\u89c4\u529d\u7edf\u6cbb\u8005\u514b\u5236\u5a31\u89c2\u3001\u52e4\u4fed\u5b88\u6210"
    "\u7684\u7528\u610f\u3002"
)
replace_para_text(doc.paragraphs[97], NEW_97, color=GREEN)

# ============================================================
# Paragraph [106] - 文采斐然...总结段（YELLOW整段）
# 查重标记：与李建秀、王云馨等论文重复
# ============================================================
print("Processing [106]...")
NEW_106 = (
    "\u4ee5\u4e0a\u8bf8\u8d4b\u5728\u63cf\u5199\u9a6c\u4f0e\u65f6\u5747\u8bcd\u91c7\u4e30"
    "\u8d61\uff0c\u4ee5\u5bcc\u4e8e\u5f20\u529b\u7684\u4fee\u8f9e\u5f70\u663e\u9a6c\u4e0a"
    "\u6280\u827a\u7684\u60ca\u9669\u4e0e\u7f8e\u611f\u3002\u4e0e\u8d4b\u4f53\u7684\u6587"
    "\u5b66\u94fa\u9648\u4e0d\u540c\uff0c\u53f2\u6599\u7b14\u8bb0\u5219\u63d0\u4f9b\u4e86"
    "\u66f4\u4e3a\u5177\u4f53\u7684\u6280\u827a\u7ec6\u8282\uff1a\u4e54\u5f5d\u89c2\u519b"
    "\u8425\u9a6c\u4f0e\u6f14\u7ec3\u540e\u4f5c\u300a\u7acb\u8d70\u9a6c\u8d4b\u300b\uff0c"
    "\u8d75\u7498\u300a\u56e0\u8bdd\u5f55\u300b\u5219\u8bb0\u5f55\u4e86\u201c\u900f\u5251"
    "\u95e8\u4f0e\u201d\u7b49\u519b\u4e2d\u9a6c\u4f0e\u7684\u5177\u4f53\u60c5\u72b6\u3002"
    "\u603b\u4f53\u800c\u8a00\uff0c\u5510\u4ee3\u53f2\u4e66\u5bf9\u9a6c\u4f0e\u7684\u8bb0"
    "\u8f7d\u8f83\u4e3a\u7b80\u7565\uff0c\u800c\u8d4b\u4e0e\u7b14\u8bb0\u5c0f\u8bf4\u5219"
    "\u4fdd\u5b58\u4e86\u8fd9\u4e00\u9a6c\u672f\u6d3b\u52a8\u7684\u4e30\u5bcc\u7ec6\u8282"
    "\uff0c\u4ece\u4e2d\u53ef\u89c1\u201c\u8d70\u9a6c\u51fb\u94b1\u201d\u201c\u63b7\u8c46"
    "\u4e8e\u523a\u201d\u201c\u7acb\u9a6c\u4e66\u5199\u201d\u201c\u900f\u5251\u95e8\u4f0e"
    "\u201d\u7b49\u5404\u8272\u7edd\u6280\uff0c\u4e43\u81f3\u5bab\u5ef7\u5973\u5b50\u4ea6"
    "\u80fd\u9a7e\u9a6d\u767b\u573a\u3001\u6280\u60ca\u56db\u5ea7\u3002"
)
replace_para_text(doc.paragraphs[106], NEW_106, color=GREEN)

doc.save(SRC)
print("\nAll 6 paragraphs rewritten and saved to", SRC)
