# -*- coding: utf-8 -*-
"""
Reduce plagiarism for the 竿技/险竿 section (paragraphs [373] and [375]).
- Paragraph [373]: rewrite the intro sentence before the poem quote
- Paragraph [375]: rewrite the analysis paragraph after the poem
- Mark rewritten text in green (008000)
- Keep poem citation [374] untouched
"""

from docx import Document
from pathlib import Path
import lxml.etree as ET
from copy import deepcopy

ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

doc_path = Path(__file__).resolve().parent.parent / \
    '\u90d1\u748e_\u300a\u6c49\u5510\u767e\u620f\u4e66\u5199\u7814\u7a76\u300b\u4e8c\u7a3f_20260226(2).docx'

doc = Document(str(doc_path))

GREEN = '008000'


def clone_pPr(source_para):
    pPr = source_para._element.find(f'{{{ns}}}pPr')
    if pPr is not None:
        return deepcopy(pPr)
    return None


def make_run(text, font='\u5b8b\u4f53', sz='24', color=None):
    r = ET.SubElement(ET.Element('dummy'), f'{{{ns}}}r')
    rPr = ET.SubElement(r, f'{{{ns}}}rPr')
    fonts_el = ET.SubElement(rPr, f'{{{ns}}}rFonts')
    fonts_el.set(f'{{{ns}}}ascii', font)
    fonts_el.set(f'{{{ns}}}hAnsi', font)
    fonts_el.set(f'{{{ns}}}hint', 'eastAsia')
    if color:
        c_el = ET.SubElement(rPr, f'{{{ns}}}color')
        c_el.set(f'{{{ns}}}val', color)
    sz_el = ET.SubElement(rPr, f'{{{ns}}}sz')
    sz_el.set(f'{{{ns}}}val', sz)
    t_el = ET.SubElement(r, f'{{{ns}}}t')
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_el.text = text
    return r


def replace_para_text(para, new_text, color=None):
    elem = para._element
    for child in list(elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'pPr':
            elem.remove(child)
    r = make_run(new_text, color=color)
    elem.append(r)


# ─── Paragraph [373]: rewrite ───
p373 = doc.paragraphs[373]
old_373 = p373.text.strip()
print(f'Old [373]: {old_373[:80]}...')

new_373 = (
    '\u7aff\u6280\u4e66\u5199\u4e2d\u4e5f\u6709\u501f\u827a\u5bd3\u7406\u7684\u8bd7\u6b4c\u3002'
    '\u7531\u4e8e\u7aff\u6280\u8868\u6f14\u7684\u5371\u9669\u6027\u6781\u9ad8\uff0c'
    '\u7a0d\u6709\u4e0d\u614e\u8868\u6f14\u8005\u4fbf\u6709\u6027\u547d\u4e4b\u865e\uff0c'
    '\u8fd9\u4e0e\u767e\u620f\u672c\u5e94\u5a31\u4eba\u7684\u521d\u8877\u5927\u76f8\u5f84\u5ead\u3002'
    '\u67f3\u66fe\u300a\u9669\u7aff\u884c\u300b\u4ee5\u89c4\u529d\u7684\u53e3\u543b\u5411'
    '\u7f18\u7aff\u827a\u4eba\u53d1\u51fa\u611f\u6168\uff0c\u5e76\u5c06\u7aff\u6280\u4e4b\u9669'
    '\u4e0e\u4ed5\u9014\u5b98\u573a\u7684\u6ce2\u8be1\u4e91\u8c32\u76f8\u7c7b\u6bd4\uff0c'
    '\u8bd7\u4e91\uff1a'
)

replace_para_text(p373, new_373, color=GREEN)
print(f'New [373]: {new_373[:80]}...')

# ─── Paragraph [375]: rewrite ───
p375 = doc.paragraphs[375]
old_375 = p375.text.strip()
print(f'\nOld [375]: {old_375[:80]}...')

new_375 = (
    '\u5168\u8bd7\u4ee5\u7f18\u7aff\u827a\u4eba\u7684\u906d\u9047\u4e3a\u5207\u5165\u70b9\uff0c'
    '\u5148\u611f\u53f9\u5176\u521d\u4ee5\u60ca\u9669\u6280\u827a\u5f15\u8d77\u5e1d\u738b\u5173\u6ce8\uff0c'
    '\u7ee7\u800c\u56e0\u8d2a\u6b32\u6539\u4e3a\u73a9\u5f04\u6743\u672f\u3001\u8c04\u5a9a\u541b\u4e0a\u3002'
    '\u67f3\u66fe\u771f\u6b63\u8981\u8868\u8fbe\u7684\u5e76\u975e\u7aff\u6280\u672c\u8eab\uff0c'
    '\u800c\u662f\u501f\u9669\u7aff\u7684\u610f\u8c61\u6620\u5c04\u5b98\u573a\u5f97\u5931\uff1a'
    '\u90a3\u4e9b\u6c89\u6eba\u4e8e\u6743\u52bf\u800c\u5931\u53bb\u541b\u6069\u8005\uff0c'
    '\u8d2c\u8c2a\u6d41\u5f99\u5929\u6daf\uff0c\u5176\u7a98\u8feb\u5904\u5883\u8f83\u4e4b\u767e\u5c3a'
    '\u9ad8\u7aff\u4e0a\u7684\u8fdb\u9000\u7ef4\u8c37\u6709\u8fc7\u4e4b\u800c\u65e0\u4e0d\u53ca\u3002'
    '\u8fd9\u5c42\u8b66\u8beb\u610f\u5473\u4f7f\u7aff\u6280\u4e66\u5199\u7a81\u7834\u4e86\u5355\u7eaf'
    '\u7684\u5947\u89c2\u63cf\u6479\uff0c\u800c\u88ab\u8d4b\u4e88\u4e86\u66f4\u6df1\u5c42\u7684'
    '\u793e\u4f1a\u8bbd\u8c15\u529f\u80fd\u3002\u6587\u5b97\u5373\u4f4d\u4ee5\u540e\uff0c'
    '\u4ee4\u884c\u7981\u6b62\u3002'
)

replace_para_text(p375, new_375, color=GREEN)
print(f'New [375]: {new_375[:80]}...')

# ─── Save ───
doc.save(str(doc_path))
print('\nDone! Saved successfully.')

# ─── Verify ───
doc2 = Document(str(doc_path))
for i in [372, 373, 374, 375, 376]:
    if i < len(doc2.paragraphs):
        p = doc2.paragraphs[i]
        t = p.text.strip()
        colors = set()
        for r in p.runs:
            rPr_el = r._element.find(f'{{{ns}}}rPr')
            if rPr_el is not None:
                c_el = rPr_el.find(f'{{{ns}}}color')
                if c_el is not None:
                    colors.add(c_el.get(f'{{{ns}}}val'))
        cstr = ','.join(colors) if colors else 'default'
        print(f'[{i}] color={cstr} | {t[:120]}')
