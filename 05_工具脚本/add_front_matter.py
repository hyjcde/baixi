"""
为《汉唐百戏书写研究》添加封面、独创性声明、英文摘要。

最终文档结构：
1. 封面（无页眉页脚）
2. 独创性声明和使用授权声明（无页眉页脚）
3. 中文摘要（已有，页眉"摘要"，罗马数字页码从I开始）
4. 英文摘要 Abstract（新增，页眉"Abstract"，罗马数字页码续编）
5. 目录（页眉"目录"，罗马数字页码续编）
6. 绪论... 正文...（阿拉伯数字页码从1开始）
"""

import copy
from docx import Document
from docx.shared import Pt, Cm, Emu, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

INPUT_FILE = "../郑璎_《汉唐百戏书写研究》_20260303.docx"
OUTPUT_FILE = "../郑璎_《汉唐百戏书写研究》_20260303.docx"

BLACK = RGBColor(0, 0, 0)


def set_run_font(run, font_name, font_size_pt, bold=False, east_asia=None):
    run.font.size = Pt(font_size_pt)
    run.font.name = font_name
    run.font.bold = bold
    run.font.color.rgb = BLACK
    if east_asia:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), east_asia)


def add_paragraph_with_text(doc, text, font_name, font_size_pt, bold=False,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_before=None, space_after=None,
                            line_spacing=None, east_asia=None,
                            insert_before_element=None):
    if insert_before_element is not None:
        new_p = parse_xml(f'<w:p {nsdecls("w")}/>')
        insert_before_element.addprevious(new_p)
        from docx.text.paragraph import Paragraph
        p = Paragraph(new_p, doc)
    else:
        p = doc.add_paragraph()

    p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, font_name, font_size_pt, bold, east_asia or font_name)

    pf = p.paragraph_format
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    if line_spacing is not None:
        pf.line_spacing = line_spacing

    return p


def create_section_break_on_paragraph(para_elem, template_sect_pr, page_num_fmt=None, start_num=None):
    """Add a section break (next page) to the given paragraph's pPr."""
    new_sect_pr = copy.deepcopy(template_sect_pr)
    for child in list(new_sect_pr):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag not in ('pgSz', 'pgMar', 'cols'):
            new_sect_pr.remove(child)

    type_elem = parse_xml(f'<w:type {nsdecls("w")} w:val="nextPage"/>')
    new_sect_pr.insert(0, type_elem)

    if page_num_fmt:
        attrs = f'w:fmt="{page_num_fmt}"'
        if start_num is not None:
            attrs += f' w:start="{start_num}"'
        pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")} {attrs}/>')
        new_sect_pr.append(pgNumType)

    pPr = para_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = parse_xml(f'<w:pPr {nsdecls("w")}/>')
        para_elem.insert(0, pPr)

    existing = pPr.find(qn('w:sectPr'))
    if existing is not None:
        pPr.remove(existing)

    pPr.append(new_sect_pr)


def set_header_content(header, text):
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run.font.color.rgb = BLACK
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="000000"/>'
        f'</w:pBdr>'
    )
    old = pPr.find(qn('w:pBdr'))
    if old is not None:
        pPr.remove(old)
    pPr.append(pBdr)


def clear_header_footer(section):
    """Make header and footer empty (for cover page and declaration)."""
    h = section.header
    h.is_linked_to_previous = False
    for p in h.paragraphs:
        p.clear()
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                pPr.remove(pBdr)

    f = section.footer
    f.is_linked_to_previous = False
    for p in f.paragraphs:
        p.clear()


def set_footer_page_number(footer):
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rpr_xml = '<w:rPr><w:color w:val="000000"/><w:sz w:val="18"/><w:szCs w:val="18"/><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/></w:rPr>'

    run_begin = parse_xml(f'<w:r {nsdecls("w")}>{rpr_xml}<w:fldChar w:fldCharType="begin"/></w:r>')
    run_instr = parse_xml(f'<w:r {nsdecls("w")}>{rpr_xml}<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>')
    run_sep = parse_xml(f'<w:r {nsdecls("w")}>{rpr_xml}<w:fldChar w:fldCharType="separate"/></w:r>')
    run_end = parse_xml(f'<w:r {nsdecls("w")}>{rpr_xml}<w:fldChar w:fldCharType="end"/></w:r>')

    p._element.append(run_begin)
    p._element.append(run_instr)
    p._element.append(run_sep)
    p._element.append(run_end)


def main():
    doc = Document(INPUT_FILE)
    body = doc.element.body

    template_sect_pr = body.find(qn('w:sectPr'))
    if template_sect_pr is None:
        print("ERROR: No section properties found")
        return
    template_sect_pr = copy.deepcopy(template_sect_pr)

    first_para = doc.paragraphs[0]._element  # "摘 要" heading

    # ========================================
    # 1. CREATE COVER PAGE
    # ========================================
    cover_paras = []

    # "索取号：" line (top left)
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="left"/></w:pPr></w:p>')
    first_para.addprevious(p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(p, doc)
    run = para.add_run("索取号：")
    set_run_font(run, '宋体', 14, east_asia='宋体')
    cover_paras.append(p)

    # Empty spacer lines
    for _ in range(3):
        sp = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
        first_para.addprevious(sp)
        cover_paras.append(sp)

    # "硕 士 学 位 论 文" title
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
    first_para.addprevious(p)
    para = Paragraph(p, doc)
    run = para.add_run("硕 士 学 位 论 文")
    set_run_font(run, '黑体', 26, bold=True, east_asia='黑体')
    cover_paras.append(p)

    # Empty spacer
    for _ in range(2):
        sp = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
        first_para.addprevious(sp)
        cover_paras.append(sp)

    # Thesis title
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
    first_para.addprevious(p)
    para = Paragraph(p, doc)
    run = para.add_run("汉唐百戏书写研究")
    set_run_font(run, '黑体', 26, bold=True, east_asia='黑体')
    cover_paras.append(p)

    # Empty spacer
    for _ in range(3):
        sp = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
        first_para.addprevious(sp)
        cover_paras.append(sp)

    # Info fields
    info_lines = [
        ("研 究 生：", "郑  璎"),
        ("指导教师：", "       （含职称）"),
        ("培养单位：", "       学院"),
        ("一级学科：", "中国语言文学"),
        ("二级学科：", "中国古代文学"),
        ("完成时间：", "2026年  月  日"),
        ("答辩时间：", "2026年  月  日"),
    ]
    for label, value in info_lines:
        p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
        first_para.addprevious(p)
        para = Paragraph(p, doc)
        run_label = para.add_run(label)
        set_run_font(run_label, '宋体', 14, bold=True, east_asia='宋体')
        run_value = para.add_run(value)
        set_run_font(run_value, '宋体', 14, east_asia='宋体')
        cover_paras.append(p)

    # Add section break after the last cover paragraph (no page number)
    last_cover = cover_paras[-1]
    create_section_break_on_paragraph(last_cover, template_sect_pr)

    # ========================================
    # 2. CREATE DECLARATION PAGE
    # ========================================
    decl_paras = []

    # Empty spacer
    sp = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
    first_para.addprevious(sp)
    decl_paras.append(sp)

    # "学 位 论 文 独 创 性 声 明"
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
    first_para.addprevious(p)
    para = Paragraph(p, doc)
    run = para.add_run("学 位 论 文 独 创 性 声 明")
    set_run_font(run, '宋体', 22, bold=True, east_asia='宋体')
    decl_paras.append(p)

    # Empty spacer
    sp = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="left"/></w:pPr></w:p>')
    first_para.addprevious(sp)
    decl_paras.append(sp)

    # Declaration text
    decl_text = "本人郑重声明：所提交的学位论文是本人在导师指导下进行的研究工作和取得的研究成果。本论文中除引文外，所有实验、数据和有关材料均是真实的。本论文中除引文和致谢的内容外，不包含其他人或其它机构已经发表或撰写过的研究成果。其他同志对本研究所做的贡献均已在论文中作了声明并表示了谢意。"
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="both"/><w:ind w:firstLineChars="200" w:firstLine="480"/></w:pPr></w:p>')
    first_para.addprevious(p)
    para = Paragraph(p, doc)
    run = para.add_run(decl_text)
    set_run_font(run, '宋体', 14, east_asia='宋体')
    para.paragraph_format.line_spacing = Pt(28)
    decl_paras.append(p)

    # Signature line
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="left"/></w:pPr></w:p>')
    first_para.addprevious(p)
    para = Paragraph(p, doc)
    run = para.add_run("学位论文作者签名：                         日      期：")
    set_run_font(run, '宋体', 14, east_asia='宋体')
    decl_paras.append(p)

    # Spacer
    for _ in range(3):
        sp = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
        first_para.addprevious(sp)
        decl_paras.append(sp)

    # "学位论文使用授权声明"
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
    first_para.addprevious(p)
    para = Paragraph(p, doc)
    run = para.add_run("学位论文使用授权声明")
    set_run_font(run, '宋体', 22, bold=True, east_asia='宋体')
    decl_paras.append(p)

    # Spacer
    sp = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="left"/></w:pPr></w:p>')
    first_para.addprevious(sp)
    decl_paras.append(sp)

    # Authorization text
    auth_text = "研究生在校攻读学位期间论文工作的知识产权单位属南京师范大学。学校有权保存本学位论文的电子和纸质文档，可以借阅或上网公布本学位论文的部分或全部内容，可以采用影印、复印等手段保存、汇编本学位论文。学校可以向国家有关机关或机构送交论文的电子和纸质文档，允许论文被查阅和借阅。"
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="both"/><w:ind w:firstLineChars="200" w:firstLine="480"/></w:pPr></w:p>')
    first_para.addprevious(p)
    para = Paragraph(p, doc)
    run = para.add_run(auth_text)
    set_run_font(run, '宋体', 14, east_asia='宋体')
    para.paragraph_format.line_spacing = Pt(28)
    decl_paras.append(p)

    # Signature lines
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="left"/></w:pPr></w:p>')
    first_para.addprevious(p)
    para = Paragraph(p, doc)
    run = para.add_run("学位论文作者签名：               指导教师签名：")
    set_run_font(run, '宋体', 14, east_asia='宋体')
    decl_paras.append(p)

    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="left"/></w:pPr></w:p>')
    first_para.addprevious(p)
    para = Paragraph(p, doc)
    run = para.add_run("日            期：               日            期：")
    set_run_font(run, '宋体', 14, east_asia='宋体')
    decl_paras.append(p)

    # Section break after declaration (no page numbers)
    last_decl = decl_paras[-1]
    create_section_break_on_paragraph(last_decl, template_sect_pr)

    # ========================================
    # 3. INSERT ENGLISH ABSTRACT after 关键词
    # ========================================
    # Find 关键词 paragraph (currently has section break to 目录)
    keywords_para = None
    keywords_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text and para.text.strip().startswith("关键词"):
            keywords_para = para
            keywords_idx = i
            break

    if keywords_para is None:
        print("ERROR: Cannot find 关键词 paragraph")
        return

    # The 关键词 paragraph currently has a section break -> 目录
    # We need to: remove that section break from 关键词, add abstract after it,
    # then put section break on the last abstract paragraph

    # Remove section break from keywords paragraph
    kw_pPr = keywords_para._element.find(qn('w:pPr'))
    kw_sectPr = kw_pPr.find(qn('w:sectPr')) if kw_pPr is not None else None
    saved_sectPr = None
    if kw_sectPr is not None:
        saved_sectPr = copy.deepcopy(kw_sectPr)
        kw_pPr.remove(kw_sectPr)

    # Now add a NEW section break on keywords para (摘要 -> Abstract transition)
    create_section_break_on_paragraph(keywords_para._element, template_sect_pr, "upperRoman")

    # Find the element AFTER keywords to insert abstract before
    # The next paragraph after keywords
    next_elem = keywords_para._element.getnext()

    # Create abstract paragraphs
    abstract_paras = []

    # "Abstract" title
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="center"/></w:pPr></w:p>')
    next_elem.addprevious(p)
    para = Paragraph(p, doc)
    run = para.add_run("Abstract")
    set_run_font(run, 'Times New Roman', 16, bold=True, east_asia='黑体')
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    abstract_paras.append(p)

    # Abstract body text (placeholder - user needs to fill in)
    abstract_text = "This study collectively terms the depictions of baixi (hundred entertainments) performance scenes, technical contents, performers, and the social evaluations and cultural discussions triggered by baixi found in Han-Tang literary sources as \"baixi writing.\" Based on Han-Tang documents, this dissertation examines the representation and evolution of baixi writing from three dimensions: literary writing, historical writing, and religious writing."
    abstract_text2 = "Chapter One clarifies and defines the concept of \"baixi\" and its related terms, proposing and defining the entirely new concept of \"baixi writing.\" Chapter Two investigates the writing of acrobatic baixi during the Han-Tang period. Chapter Three examines the writing of animal-taming and animal-imitation baixi. Chapter Four explores the writing of magic and illusion baixi."

    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="both"/><w:ind w:firstLineChars="200" w:firstLine="480"/></w:pPr></w:p>')
    next_elem.addprevious(p)
    para = Paragraph(p, doc)
    run = para.add_run(abstract_text)
    set_run_font(run, 'Times New Roman', 12, east_asia='宋体')
    para.paragraph_format.line_spacing = Pt(20)
    abstract_paras.append(p)

    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="both"/><w:ind w:firstLineChars="200" w:firstLine="480"/></w:pPr></w:p>')
    next_elem.addprevious(p)
    para = Paragraph(p, doc)
    run = para.add_run(abstract_text2)
    set_run_font(run, 'Times New Roman', 12, east_asia='宋体')
    para.paragraph_format.line_spacing = Pt(20)
    abstract_paras.append(p)

    # Empty line
    sp = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="left"/></w:pPr></w:p>')
    next_elem.addprevious(sp)
    abstract_paras.append(sp)

    # Key words
    p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:jc w:val="left"/></w:pPr></w:p>')
    next_elem.addprevious(p)
    para = Paragraph(p, doc)
    run_label = para.add_run("Key words: ")
    set_run_font(run_label, 'Times New Roman', 12, bold=True, east_asia='宋体')
    run_value = para.add_run("baixi; Han-Tang; acrobatics; animal-taming; animal-imitation; magic and illusion; literary writing; historical writing")
    set_run_font(run_value, 'Times New Roman', 12, east_asia='宋体')
    abstract_paras.append(p)

    # Add section break on last abstract paragraph (Abstract -> 目录)
    last_abstract = abstract_paras[-1]
    create_section_break_on_paragraph(last_abstract, template_sect_pr, "upperRoman")

    # ========================================
    # 4. SAVE AND RELOAD TO FIX HEADERS
    # ========================================
    doc.save(OUTPUT_FILE)
    print("Saved intermediate document with front matter content")

    # Reload to set headers/footers for new sections
    doc = Document(OUTPUT_FILE)
    print(f"Document now has {len(doc.sections)} sections")

    # Expected sections after insertion:
    # 0: Cover page (no header/footer)
    # 1: Declaration page (no header/footer)
    # 2: 摘要 (header="摘要", Roman I start)
    # 3: Abstract (header="Abstract", Roman continued)
    # 4: 目录 (header="目录", Roman continued)
    # 5: 绪论 (header="绪论", decimal 1 start)
    # 6: 第一章
    # 7: 第二章
    # 8: 第三章
    # 9: 第四章
    # 10: 结语
    # 11: 附录 (附录1)
    # 12: 附录 (附录2)
    # 13: 参考文献

    section_configs = [
        {"header": None,        "footer": False},  # Cover
        {"header": None,        "footer": False},  # Declaration
        {"header": "摘要",      "footer": True},   # 摘要
        {"header": "Abstract",  "footer": True},   # Abstract
        {"header": "目录",      "footer": True},   # 目录
        {"header": "绪论",      "footer": True},   # 绪论
        {"header": "第一章 汉唐百戏书写概述",    "footer": True},
        {"header": "第二章 杂技类百戏书写研究",  "footer": True},
        {"header": "第三章 驯兽与拟兽类百戏书写研究", "footer": True},
        {"header": "第四章 幻术类百戏书写研究",  "footer": True},
        {"header": "结语",      "footer": True},
        {"header": "附录",      "footer": True},
        {"header": "附录",      "footer": True},
        {"header": "参考文献",  "footer": True},
    ]

    num_sections = len(doc.sections)
    print(f"Expected {len(section_configs)} sections, got {num_sections}")

    for i, section in enumerate(doc.sections):
        if i >= len(section_configs):
            break
        cfg = section_configs[i]

        if cfg["header"] is None:
            clear_header_footer(section)
        else:
            set_header_content(section.header, cfg["header"])
            if cfg["footer"]:
                set_footer_page_number(section.footer)

        print(f"Section {i}: header={'(none)' if cfg['header'] is None else cfg['header']}")

    doc.save(OUTPUT_FILE)
    print(f"\nFinal document saved to {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
