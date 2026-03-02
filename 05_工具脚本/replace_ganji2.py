# -*- coding: utf-8 -*-
"""
Reduce plagiarism for the 竿技 section part 2:
paragraphs [361], [363], [364], [366], [367], [369]
Mark rewritten text in green (008000).
"""

from docx import Document
from pathlib import Path
import lxml.etree as ET
from copy import deepcopy

ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

doc_path = Path(__file__).resolve().parent.parent / \
    '\u90d1\u748e_\u300a\u6c49\u5510\u767e\u620f\u4e66\u5199\u7814\u7a76\u300b\u4e8c\u7a3f_20260226(2).docx'

doc = Document(str(doc_path))

GREEN = '008000'


def make_run(text, font='\u5b8b\u4f53', sz='24', color=None):
    r = ET.SubElement(ET.Element('dummy'), f'{{{ns}}}r')
    rPr = ET.SubElement(r, f'{{{ns}}}rPr')
    fonts_el = ET.SubElement(rPr, f'{{{ns}}}rFonts')
    fonts_el.set(f'{{{ns}}}ascii', font)
    fonts_el.set(f'{{{ns}}}hAnsi', font)
    fonts_el.set(f'{{{ns}}}hint', 'eastAsia')
    if color:
        c_el = ET.SubElement(rPr, f'{{{ns}}}color')
        c_el.set(f'{{{ns}}}val', color)
    sz_el = ET.SubElement(rPr, f'{{{ns}}}sz')
    sz_el.set(f'{{{ns}}}val', sz)
    t_el = ET.SubElement(r, f'{{{ns}}}t')
    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_el.text = text
    return r


def replace_para_text(para, new_text, color=None):
    elem = para._element
    for child in list(elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'pPr':
            elem.remove(child)
    r = make_run(new_text, color=color)
    elem.append(r)


def replace_para_mixed(para, segments):
    """Replace paragraph with mixed color segments.
    segments: list of (text, color_or_None)
    """
    elem = para._element
    for child in list(elem):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'pPr':
            elem.remove(child)
    for text, color in segments:
        r = make_run(text, color=color)
        elem.append(r)


# ─── [361] ───
p361 = doc.paragraphs[361]
print(f'Old [361]: {p361.text.strip()}')
# "又《全唐诗》中收一歌谣，曰《幽州谣》：" → rewrite
new_361 = (
    '\u300a\u5e7d\u5dde\u8c23\u300b\u4ee5\u6b4c\u8c23\u5f62\u5f0f\uff0c'
    '\u5c06\u7aff\u6280\u4e0e\u65f6\u4ee3\u4e71\u4ea1\u76f8\u8054\u7cfb\uff0c'
    '\u4ece\u5355\u7eaf\u6280\u827a\u63cf\u5199\u8f6c\u5411\u5386\u53f2\u611f\u6168\uff1a'
)
replace_para_text(p361, new_361, color=GREEN)
print(f'New [361]: {new_361}')

# ─── [363] ───
# Structure: intro sentence (keep) + long quotation (keep) + analysis (rewrite)
# The quotation ends at "……契丹初闻莫悟，至是而应之。"
# After that is the analysis to rewrite
p363 = doc.paragraphs[363]
old_363 = p363.text.strip()
print(f'\nOld [363]: {old_363[:100]}...')

# Split: keep the intro + quotation, rewrite the analysis
# Quotation ends with: 契丹初闻莫悟，至是而应之。"
split_marker = '\u5951\u4e39\u521d\u95fb\u83ab\u609f\uff0c\u81f3\u662f\u800c\u5e94\u4e4b\u3002\u201d'
idx = old_363.find(split_marker)
if idx >= 0:
    keep_part = old_363[:idx + len(split_marker)]
    old_analysis = old_363[idx + len(split_marker):]
    print(f'  Keep: ...{keep_part[-60:]}')
    print(f'  Old analysis: {old_analysis}')
else:
    keep_part = old_363
    old_analysis = ''
    print('  WARNING: split marker not found!')

new_analysis = (
    '\u636e\u6b64\u53ef\u77e5\uff0c\u8fd9\u6279\u6234\u7aff\u827a\u4eba\u539f\u7cfb'
    '\u7384\u5b97\u8d50\u4e88\u5b89\u7984\u5c71\u7684\u5bab\u5ef7\u4f0e\u4eba\uff0c'
    '\u8f97\u8f6c\u4f20\u4e60\u540e\u5f97\u4e94\u767e\u4f59\u4eba\uff0c'
    '\u6280\u827a\u7cbe\u6e5b\u3002\u5948\u4f55\u5f53\u595a\u3001\u5951\u4e39'
    '\u4e24\u8fb9\u6765\u72af\u65f6\uff0c\u5b88\u5c06\u65e0\u5175\u53ef\u7528\uff0c'
    '\u7adf\u5c06\u8fd9\u4e9b\u8eab\u6000\u7edd\u6280\u7684\u7aff\u6280\u827a\u4eba'
    '\u62bd\u8c03\u4e0a\u9635\uff0c\u7ed3\u679c\u5728\u6e05\u6c34\u6cb3\u7554\u5927\u8d25\uff0c'
    '\u51e0\u4e4e\u5168\u519b\u8986\u6ca1\uff0c\u4ec5\u4e09\u4eba\u4f0f\u8349\u800c\u514d\u3002'
    '\u76f8\u4f20\u5728\u654c\u519b\u62b5\u8fbe\u524d\u4e00\u4e2a\u6708\uff0c'
    '\u5e7d\u5dde\u6c11\u95f4\u5df2\u6709\u6b64\u8c23\u4f20\u5531\uff0c'
    '\u65f6\u4eba\u4e0d\u89e3\u5176\u610f\uff0c\u4e8b\u540e\u65b9\u77e5\u7adf\u662f'
    '\u4e00\u8bed\u6210\u8c36\u3002'
)

replace_para_mixed(p363, [
    (keep_part, None),
    (new_analysis, GREEN),
])
print(f'  New analysis: {new_analysis}')

# ─── [364] ───
p364 = doc.paragraphs[364]
print(f'\nOld [364]: {p364.text.strip()}')
new_364 = (
    '\u4e2d\u665a\u5510\u4ee5\u964d\uff0c\u4eba\u4eec\u5bf9\u7aff\u6728\u6280\u5de7\u7684'
    '\u5ba1\u7f8e\u5df2\u4e0d\u6b62\u4e8e\u96be\u5ea6\u4e0e\u82b1\u6837\uff0c'
    '\u8f6c\u800c\u8ffd\u6c42\u66f4\u4e3a\u60ca\u5fc3\u52a8\u9b44\u7684\u89c6\u89c9\u523a\u6fc0\u3002'
    '\u82cf\u9e57\u300a\u675c\u9633\u6742\u7f16\u300b\u4e2d\u7559\u4e0b\u4e86'
    '\u656c\u5b97\u671d\u5343\u79cb\u8282\u7aff\u6728\u8868\u6f14\u7684\u4e00\u6bb5'
    '\u751f\u52a8\u8bb0\u8f7d\uff1a'
)
replace_para_text(p364, new_364, color=GREEN)
print(f'New [364]: {new_364}')

# ─── [366] ───
p366 = doc.paragraphs[366]
print(f'\nOld [366]: {p366.text.strip()}')
new_366 = (
    '\u7aff\u6280\u4e0e\u97f3\u4e50\u914d\u5408\uff0c'
    '\u827a\u4eba\u4ef0\u4fef\u7ffb\u98de\u4e4b\u95f4\u8e0f\u51c6\u97f3\u4e50\u8282\u62cd\uff0c'
    '\u6280\u672f\u4e4b\u9ad8\u8d85\u3001\u573a\u9762\u4e4b\u60ca\u9669\uff0c'
    '\u5747\u53ef\u4ece\u4e0b\u9762\u4e24\u9996\u8bd7\u4f5c\u4e2d\u5f97\u5230\u5370\u8bc1\u3002'
)
replace_para_text(p366, new_366, color=GREEN)
print(f'New [366]: {new_366}')

# ─── [367] ───
p367 = doc.paragraphs[367]
print(f'\nOld [367]: {p367.text.strip()}')
new_367 = '\u987e\u51b5\u300a\u9669\u7aff\u6b4c\u300b\u4e91\uff1a'
replace_para_text(p367, new_367, color=GREEN)
print(f'New [367]: {new_367}')

# ─── [369] ───
p369 = doc.paragraphs[369]
print(f'\nOld [369]: {p369.text.strip()[:100]}...')
new_369 = (
    '\u6b64\u8bd7\u5927\u7ea6\u4f5c\u4e8e\u987e\u51b5\u4efb\u9547\u6d77\u519b'
    '\u8282\u5ea6\u4f7f\u5224\u5b98\u671f\u95f4\u3002\u5b9b\u9675\uff0c'
    '\u5373\u968b\u4ee3\u6539\u540d\u4e4b\u5ba3\u57ce\uff0c\u5510\u4ee3\u96b6\u5c5e'
    '\u5ba3\u5dde\u9547\u6d77\u519b\u7ba1\u8f96\u3002\u8bd7\u4e2d\u523b\u753b\u7684'
    '\u5b9b\u9675\u5973\u5b50\u5728\u60ac\u7a7a\u6a2a\u7aff\u4e0a\u884c\u8d70\u81ea\u5982\uff0c'
    '\u5982\u5c65\u5e73\u5730\uff0c\u7ffb\u8eab\u6302\u5f71\u3001\u76d8\u65cb\u4f3c\u98ce\uff0c'
    '\u5176\u8eab\u624b\u654f\u6377\u5ea6\u8d85\u8d8a\u98de\u9e1f\u3001\u60ca\u733f\uff0c'
    '\u5c3d\u663e\u5973\u6027\u827a\u4eba\u7684\u80c6\u9b44\u4e0e\u7075\u5de7\u3002'
    '\u5728\u6beb\u65e0\u501f\u529b\u7684\u957f\u7aff\u4e4b\u4e0a\uff0c'
    '\u8868\u6f14\u8005\u65e2\u80fd\u8eab\u8f7b\u5982\u71d5\u3001\u4e0a\u4e0b\u7ffb\u98de\uff0c'
    '\u53c8\u80fd\u5012\u6302\u3001\u8170\u65cb\u3001\u76d8\u8e4b\uff0c'
    '\u4e00\u5957\u52a8\u4f5c\u4e00\u6c14\u5475\u6210\u3001\u75be\u5982\u95ea\u7535\uff0c'
    '\u4ee4\u89c2\u8005\u53f9\u4e3a\u89c2\u6b62\u3002'
)
replace_para_text(p369, new_369, color=GREEN)
print(f'New [369]: {new_369[:80]}...')

# ─── Save ───
doc.save(str(doc_path))
print('\n\nDone! Saved successfully.')

# ─── Verify ───
doc2 = Document(str(doc_path))
for i in [361, 362, 363, 364, 365, 366, 367, 368, 369]:
    if i < len(doc2.paragraphs):
        p = doc2.paragraphs[i]
        t = p.text.strip()
        colors = set()
        for r in p.runs:
            rPr_el = r._element.find(f'{{{ns}}}rPr')
            if rPr_el is not None:
                c_el = rPr_el.find(f'{{{ns}}}color')
                if c_el is not None:
                    colors.add(c_el.get(f'{{{ns}}}val'))
        cstr = ','.join(colors) if colors else 'default'
        print(f'[{i}] color={cstr} | {t[:120]}')
