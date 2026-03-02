# -*- coding: utf-8 -*-
"""
Fix typos in green paragraph [143]:
1. 萌蕤 → 萌蘖
2. 讲戈之礼 → 讲戎之礼
3. 戏谐 → 戏谑 (2 occurrences)
4. 递嘿 → 递嬗
"""

from docx import Document
from pathlib import Path
import lxml.etree as ET

ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

doc_path = Path(__file__).resolve().parent.parent / \
    '\u90d1\u748e_\u300a\u6c49\u5510\u767e\u620f\u4e66\u5199\u7814\u7a76\u300b\u4e8c\u7a3f_20260226(2).docx'

doc = Document(str(doc_path))

p143 = doc.paragraphs[143]

fixes = [
    ('\u840c\u8564', '\u840c\u8616'),    # 萌蕤 → 萌蘖
    ('\u8bb2\u6208\u4e4b\u793c', '\u8bb2\u620e\u4e4b\u793c'),  # 讲戈之礼 → 讲戎之礼
    ('\u620f\u8c10', '\u620f\u8c11'),    # 戏谐 → 戏谑
    ('\u9012\u563f', '\u9012\u5b17'),    # 递嘿 → 递嬗
]

for run in p143.runs:
    if run.text:
        original = run.text
        for old, new in fixes:
            if old in run.text:
                run.text = run.text.replace(old, new)
                print(f'Fixed in run: "{old}" -> "{new}"')

doc.save(str(doc_path))
print('\nDone! Saved.')

# Verify
doc2 = Document(str(doc_path))
text = doc2.paragraphs[143].text
for old, new in fixes:
    if old in text:
        print(f'WARNING: "{old}" still present!')
    if new in text:
        print(f'OK: "{new}" found.')
